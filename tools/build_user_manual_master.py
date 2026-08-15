import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from sections.cover_and_front import add_cover_page, add_table_of_contents, add_list_of_figures
from sections.sec1_to_5 import add_sections_1_to_5
from sections.sec6_functionalities import add_section_6
from sections.sec7_to_12 import add_sections_7_to_12

# Calibrated Table of Contents
FINAL_TOC = [
    ("1. Project Overview", 4, 1),
    ("1.1 Purpose of the Project", 4, 2),
    ("1.2 Main Objectives", 4, 2),
    ("1.3 Target Users and Stakeholders", 5, 2),
    ("1.4 Major Functions of the System", 5, 2),
    ("2. System Requirements", 5, 1),
    ("2.1 Hardware Requirements", 5, 2),
    ("2.1.1 Client Smartphone Requirements", 5, 3),
    ("2.1.2 Administrator Workstation Requirements", 6, 3),
    ("2.2 Software Requirements", 6, 2),
    ("2.2.1 Application & Runtime Environment", 6, 3),
    ("2.2.2 Backend, Database & Cloud Services", 6, 3),
    ("3. Installation and Setup", 6, 1),
    ("3.1 Mobile Application Installation (Rider & Owner)", 6, 2),
    ("3.2 Web Administrative Portal Access", 7, 2),
    ("3.3 Initial System Configuration", 7, 2),
    ("4. Getting Started", 7, 1),
    ("4.1 Launching the Application", 7, 2),
    ("4.2 Role Selection & Onboarding", 7, 2),
    ("4.3 Registration and Authentication", 8, 2),
    ("4.4 Navigation & Bottom Bar Overview", 8, 2),
    ("5. User Roles and Access Permissions", 8, 1),
    ("5.1 System Administrator", 8, 2),
    ("5.2 Parking Space Owner (Host)", 9, 2),
    ("5.3 Rider / Vehicle Driver", 9, 2),
    ("5.4 Access Control Matrix", 9, 2),
    ("6. Functionalities and Usage Instructions", 11, 1),
    ("6.1 Onboarding & KYC Identity Verification", 11, 2),
    ("6.1.1 Role Selection Workflow", 11, 3),
    ("6.1.2 User Registration & Account Creation", 12, 3),
    ("6.1.3 Rider KYC Identity Verification", 13, 3),
    ("6.1.4 Space Owner KYC & Property Ownership Verification", 14, 3),
    ("6.2 Rider Module Operations", 15, 2),
    ("6.2.1 Interactive Map & Real-Time Parking Discovery", 15, 3),
    ("6.2.2 Parking Spot Listings & Filter Criteria", 16, 3),
    ("6.2.3 AI-Powered Smart Parking Natural Language Assistant", 17, 3),
    ("6.2.4 AI Smart Recommendations & Decision Rationales", 18, 3),
    ("6.2.5 Vehicle Management (My Garage)", 19, 3),
    ("6.2.6 Adding a New Vehicle to Garage", 20, 3),
    ("6.2.7 Digital Wallet & Top-up via SSLCommerz", 21, 3),
    ("6.2.8 Active Parking Session & Live Countdown Timer", 22, 3),
    ("6.2.9 Rider Booking History & Receipts", 23, 3),
    ("6.2.10 Rider Push Notifications & Event Alerts", 24, 3),
    ("6.2.11 Rider Profile & Account Settings", 25, 3),
    ("6.3 Parking Space Owner Module Operations", 26, 2),
    ("6.3.1 Space Owner Dashboard & Earnings Summary", 26, 3),
    ("6.3.2 Adding a Parking Spot – Photos & Verification Video", 27, 3),
    ("6.3.3 Adding a Parking Spot – Vehicle Slots, Pricing & Amenities", 28, 3),
    ("6.3.4 Adding a Parking Spot – Booking Mode & Availability Schedule", 29, 3),
    ("6.3.5 Adding a Parking Spot – Map Location Pinning & Publishing", 30, 3),
    ("6.3.6 Managing Parking Listings & Slot Availability", 31, 3),
    ("6.3.7 Spot QR Code Generation & Physical Standee", 32, 3),
    ("6.3.8 Owner Booking Management & Driver Tracking", 33, 3),
    ("6.3.9 Owner Wallet & Bank Withdrawal", 34, 3),
    ("6.3.10 Owner Notifications & Overstay Alerts", 35, 3),
    ("6.3.11 Owner Profile & KYC Status", 36, 3),
    ("6.4 System Administrator Module Operations", 37, 2),
    ("6.4.1 Admin Portal Authentication", 37, 3),
    ("6.4.2 Admin Executive Dashboard & Financial Overview", 38, 3),
    ("6.4.3 User Account Management & Moderation", 39, 3),
    ("6.4.4 KYC Document Review & Verification", 40, 3),
    ("6.4.5 Parking Listings Supervision & Moderation", 41, 3),
    ("6.4.6 Live Booking Monitoring & Overstay Supervision", 42, 3),
    ("6.4.7 Payment Monitoring & Financial Ledger", 43, 3),
    ("6.4.8 Revenue Reports & Analytics", 44, 3),
    ("6.4.9 Platform Configuration & Dynamic Pricing Policies", 45, 3),
    ("7. Input and Output Description", 46, 1),
    ("7.1 Rider Module Inputs and Outputs", 46, 2),
    ("7.2 Parking Space Owner Module Inputs and Outputs", 47, 2),
    ("7.3 System Administrator Module Inputs and Outputs", 48, 2),
    ("8. Error Messages and Troubleshooting", 49, 1),
    ("9. Frequently Asked Questions (FAQ)", 51, 1),
    ("9.1 General Platform Questions", 51, 2),
    ("9.2 Rider / Vehicle Driver Questions", 51, 2),
    ("9.3 Parking Space Owner Questions", 52, 2),
    ("10. Logout and Exit Procedure", 53, 1),
    ("10.1 Mobile Application Logout", 53, 2),
    ("10.2 Admin Web Portal Logout", 53, 2),
    ("11. Limitations and Precautions", 53, 1),
    ("12. Contact and Support Information", 54, 1),
    ("Appendix", 55, 1),
    ("Appendix A: System Architecture & Technical Specifications", 55, 2),
    ("Appendix B: Database Entity Summary", 55, 2),
    ("Supervisor Approval", 56, 1)
]

# Calibrated List of Figures
FINAL_LOF = [
    ("6.1", "Role Selection Screen", 11),
    ("6.2", "User Registration and Account Creation Screen", 12),
    ("6.3", "Rider KYC Identity Verification Screen", 13),
    ("6.4", "Parking Space Owner KYC Verification Screen", 14),
    ("6.5", "Interactive Map and Nearby Parking Discovery Screen", 15),
    ("6.6", "Parking Spot Listings Screen", 16),
    ("6.7", "AI-Powered Smart Parking Assistant Screen", 17),
    ("6.8", "AI Smart Recommendations Result Screen", 18),
    ("6.9", "Rider Garage (My Vehicles) Management Screen", 19),
    ("6.10", "Add New Vehicle Screen", 20),
    ("6.11", "Digital Wallet and Top-up Screen", 21),
    ("6.12", "Active Parking Session and Live Countdown Timer Screen", 22),
    ("6.13", "Rider Booking History and Records Screen", 23),
    ("6.14", "Rider Push Notifications Screen", 24),
    ("6.15", "Rider Profile and Account Settings Screen", 25),
    ("6.16", "Space Owner Dashboard and Earnings Summary Screen", 26),
    ("6.17", "Add Parking Spot – Photos and Video Upload Screen", 27),
    ("6.18", "Add Parking Spot – Vehicle Slots, Pricing, and Amenities Screen", 28),
    ("6.19", "Add Parking Spot – Booking Mode and Availability Schedule Screen", 29),
    ("6.20", "Add Parking Spot – Map Location Pinning and Publishing Screen", 30),
    ("6.21", "My Listings and Slot Capacity Screen", 31),
    ("6.22", "Parking Spot QR Code and Check-in Standee Screen", 32),
    ("6.23", "Owner Booking Management and Driver Tracking Screen", 33),
    ("6.24", "Owner Wallet and Bank Withdrawal Screen", 34),
    ("6.25", "Owner Notifications and Overstay Alerts Screen", 35),
    ("6.26", "Owner Profile and Verification Status Screen", 36),
    ("6.27", "Admin Portal Authentication Screen", 37),
    ("6.28", "Admin Executive Dashboard and Financial Overview Screen", 38),
    ("6.29", "User Account Management and Moderation Screen", 39),
    ("6.30", "KYC Document Review and Verification Screen", 40),
    ("6.31", "Parking Listings Supervision and Moderation Screen", 41),
    ("6.32", "Live Booking Monitoring and Overstay Supervision Screen", 42),
    ("6.33", "Payment Monitoring and Financial Ledger Screen", 43),
    ("6.34", "Revenue Reports and Analytics Screen", 44),
    ("6.35", "Platform Configuration and Dynamic Pricing Policies Screen", 45)
]

def generate_manual_docx(output_path, toc_data=FINAL_TOC, lof_data=FINAL_LOF):
    doc = Document()
    
    # ---------------- PAGE SETUP ----------------
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.different_first_page_header_footer = True
    
    # ---------------- DEFAULT BODY STYLE ----------------
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(11)
    font_normal.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # ---------------- RUNNING HEADER ----------------
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "ParkFinity – SPL-2 Final Project User Manual | IIT, NSTU"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.name = "Times New Roman"
    hp.runs[0].font.size = Pt(8.5)
    hp.runs[0].font.italic = True
    hp.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    
    # ---------------- RUNNING FOOTER ----------------
    footer = section.footer
    tbl_footer = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    tbl_footer.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_f_left, c_f_right = tbl_footer.cell(0, 0), tbl_footer.cell(0, 1)
    c_f_left.width = Inches(4.5)
    c_f_right.width = Inches(2.0)
    
    p_fl = c_f_left.paragraphs[0]
    p_fl.text = "Software Project Lab-II (SE 3112)"
    p_fl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_fl.runs[0].font.name = "Times New Roman"
    p_fl.runs[0].font.size = Pt(9)
    p_fl.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    p_fr = c_f_right.paragraphs[0]
    p_fr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_pg = p_fr.add_run("Page ")
    r_pg.font.name = "Times New Roman"
    r_pg.font.size = Pt(9)
    r_pg.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    fld_page = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w'))
    p_fr.runs[0]._r.append(fld_page)
    
    r_of = p_fr.add_run(" of ")
    r_of.font.name = "Times New Roman"
    r_of.font.size = Pt(9)
    r_of.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    fld_numpages = parse_xml(r'<w:fldSimple %s w:instr="NUMPAGES"/>' % nsdecls('w'))
    r_of._r.append(fld_numpages)
    
    # 1. Cover Page
    print("Generating Cover Page...")
    add_cover_page(doc)
    
    # 2. Table of Contents
    print("Generating Table of Contents...")
    add_table_of_contents(doc, toc_data)
    
    # 3. List of Figures
    print("Generating List of Figures...")
    add_list_of_figures(doc, lof_data)
    
    # 4. Sections 1 to 5
    print("Generating Sections 1 to 5...")
    add_sections_1_to_5(doc)
    
    # 5. Section 6 (Functionalities with all 35 Figures)
    print("Generating Section 6 (All 35 Screenshots & Instructions)...")
    add_section_6(doc)
    
    # 6. Sections 7 to 12 & Supervisor Approval
    print("Generating Sections 7 to 12, Appendix, and Approval Page...")
    add_sections_7_to_12(doc)
    
    # Save document
    doc.save(output_path)
    print(f"\nDocument generated successfully at: {output_path}")

if __name__ == "__main__":
    out_file = os.path.abspath("ParkFinity_User_Manual_Report.docx")
    generate_manual_docx(out_file)
