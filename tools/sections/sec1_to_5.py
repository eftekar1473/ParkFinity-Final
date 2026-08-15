import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from .common import (
    COLOR_PRIMARY, COLOR_SECONDARY, COLOR_DARK, COLOR_MUTED,
    add_heading_1, add_heading_2, add_heading_3,
    add_body_p, add_bullet, add_callout, format_styled_table
)

def add_sections_1_to_5(doc):
    # =========================================================================
    # 1. PROJECT OVERVIEW
    # =========================================================================
    add_heading_1(doc, "1. Project Overview")
    
    add_body_p(doc, 
        "With rapid urbanization and exponential growth in vehicular density across major metropolitan cities in "
        "Bangladesh—such as Dhaka, Chattogram, Sylhet, and regional academic/economic hubs—finding safe, legal, and "
        "affordable parking spaces has become one of the most stressful daily challenges for vehicle owners. Drivers "
        "routinely waste 15 to 30 minutes circling city blocks, resulting in massive productivity losses, heightened fuel "
        "consumption, increased carbon emissions, and severe roadside gridlocks. Concurrently, thousands of private parking "
        "spaces, residential building garages, commercial basements, and vacant secure plots remain completely empty or "
        "underutilized during work and business hours without any accessible medium to monetize them."
    )
    
    add_body_p(doc,
        "ParkFinity is a state-of-the-art, smart, cloud-native parking space optimization ecosystem designed to solve this "
        "severe urban imbalance. Operating as a dual-sided peer-to-peer (P2P) and commercial marketplace, ParkFinity bridges "
        "the gap between drivers in urgent need of parking (Riders) and property owners with available space (Hosts). "
        "Through an intuitive Flutter mobile application and a high-performance web administration portal, ParkFinity "
        "automates real-time geospatial parking discovery, dynamic pricing calculations, automated slot reservation, "
        "contactless digital wallet transactions, AI-powered smart recommendations, and QR-code-based check-in/check-out "
        "validation."
    )
    
    add_heading_2(doc, "1.1 Purpose of the Project")
    add_body_p(doc,
        "The primary purpose of ParkFinity is to transform urban parking management from a chaotic, manual searching "
        "chore into an organized, automated, and frictionless digital experience. By maximizing the utility of existing "
        "underutilized private and commercial parking assets, ParkFinity reduces illegal street parking, alleviates traffic "
        "bottlenecks, enables property hosts to generate dependable passive revenue, and provides vehicle drivers with "
        "guaranteed, safe, and stress-free parking reservations."
    )
    
    add_heading_2(doc, "1.2 Main Objectives")
    add_bullet(doc, "Geospatial Real-Time Discovery: ", 
        "Enable drivers to locate, evaluate, and navigate to nearby available parking spaces using interactive Google Maps with live availability status.")
    add_bullet(doc, "Peer-to-Peer Space Monetization: ", 
        "Empower residential homeowners, commercial entities, and parking lot operators to easily list, configure, and monetize their unused parking capacity.")
    add_bullet(doc, "AI-Driven Intelligent Matching: ", 
        "Incorporate an AI recommendation engine that analyzes natural-language driver queries, vehicle dimensions, proximity, pricing tiers, and amenities to offer optimal parking suggestions.")
    add_bullet(doc, "Robust KYC Identity Verification: ", 
        "Enforce rigorous verification of National Identity Cards (NID), driving licenses, and property ownership deeds to ensure maximum security, mutual trust, and fraud prevention.")
    add_bullet(doc, "Frictionless Contactless Payments: ", 
        "Integrate a secure in-app digital wallet powered by the SSLCommerz payment gateway, supporting bKash, Nagad, debit/credit cards, and automated escrow-style transaction settlements.")
    add_bullet(doc, "Automated Check-in & Overstay Prevention: ", 
        "Implement dedicated QR-code scanning at parking locations to validate driver arrivals, initiate live parking countdowns, prevent slot overstaying, and enforce automated penalty rules.")
    add_bullet(doc, "Centralized Administrative Governance: ", 
        "Provide platform administrators with comprehensive real-time dashboards for user moderation, KYC vetting, listing moderation, financial auditing, dynamic commission control, and analytics.")

    add_heading_2(doc, "1.3 Target Users and Stakeholders")
    add_bullet(doc, "Riders (Vehicle Drivers): ", 
        "Commuters, commercial car drivers, motorcyclists, SUV owners, and delivery personnel seeking instant or scheduled parking spaces.")
    add_bullet(doc, "Parking Space Owners (Hosts): ", 
        "Homeowners with vacant driveways, residential building management committees, shopping mall operators, commercial garage owners, and educational/private plot owners.")
    add_bullet(doc, "System Administrators: ", 
        "Operations managers, customer support agents, compliance vetting officers, and financial analysts responsible for overseeing platform health, legal compliance, and monetary payouts.")

    add_heading_2(doc, "1.4 Major Functions of the System")
    add_body_p(doc,
        "The ParkFinity architecture encompasses five core functional subsystems working collaboratively across client mobile "
        "apps, administrative web interfaces, and cloud database engines:"
    )
    add_bullet(doc, "Geospatial Search & Filtering: ", "Interactive map clustering, radius search, vehicle-type filtering (Car, Motorcycle, Pickup), and amenity filters (CCTV, EV charging, covered parking, 24/7 security).")
    add_bullet(doc, "Smart AI Assistant: ", "Natural-language query processing allowing riders to specify constraints like 'cheapest covered spot near campus' and receive ranked recommendations.")
    add_bullet(doc, "Vehicle Garage Registry: ", "Multi-vehicle management allowing riders to save car/bike details, license plates, and vehicle classes for one-touch booking.")
    add_bullet(doc, "Multi-Tier Dynamic Pricing & Multipliers: ", "Flexible hourly, daily, weekly, and monthly rates with platform-configured peak-hour multipliers and automated refund policies.")
    add_bullet(doc, "QR Verification & Live Session Management: ", "Physical QR stands mounted at spots enabling instant check-in, real-time remaining-time countdown, session extension, and automatic slot release.")

    # =========================================================================
    # 2. SYSTEM REQUIREMENTS
    # =========================================================================
    add_heading_1(doc, "2. System Requirements")
    add_body_p(doc,
        "To ensure high responsiveness, real-time location precision, and smooth graphical rendering, the ParkFinity "
        "ecosystem requires standard client hardware, server infrastructure, and software dependencies."
    )
    
    add_heading_2(doc, "2.1 Hardware Requirements")
    
    add_heading_3(doc, "2.1.1 Client Smartphone Requirements (Riders and Space Owners)")
    add_bullet(doc, "Operating System: ", "Android OS 8.0 (Oreo / API Level 26) or higher, or iOS 13.0 or later.")
    add_bullet(doc, "Processor & Architecture: ", "Quad-Core 1.8 GHz ARM64 processor (Octa-Core recommended).")
    add_bullet(doc, "RAM: ", "Minimum 2 GB RAM (3 GB or higher recommended for smooth map rendering).")
    add_bullet(doc, "Storage: ", "Minimum 150 MB free internal flash storage for app binary, map caching, and document buffers.")
    add_bullet(doc, "Sensors & Hardware: ", "Built-in GPS / GLONASS location sensor, rear autofocus camera (minimum 5 MP) for QR scanning and KYC document photography.")
    add_bullet(doc, "Network Connectivity: ", "Active 3G / 4G / 5G Mobile Data or Wi-Fi internet connection (minimum 1 Mbps bandwidth).")
    
    add_heading_3(doc, "2.1.2 Administrator Workstation Requirements")
    add_bullet(doc, "Computer System: ", "Modern desktop or laptop running Windows 10/11, macOS 11+, or Ubuntu Linux 20.04+.")
    add_bullet(doc, "Processor: ", "Dual-Core 2.0 GHz x86-64 / Apple Silicon processor (Intel Core i3 / AMD Ryzen 3 or higher).")
    add_bullet(doc, "System Memory (RAM): ", "Minimum 4 GB RAM (8 GB recommended for multi-tab operational monitoring).")
    add_bullet(doc, "Display Resolution: ", "1366 × 768 pixels minimum (1920 × 1080 Full HD recommended for dashboard charts).")
    add_bullet(doc, "Internet Connection: ", "Broadband connection with minimum 5 Mbps download/upload speed.")

    add_heading_2(doc, "2.2 Software Requirements")
    
    add_heading_3(doc, "2.2.1 Application & Runtime Environment")
    add_bullet(doc, "Mobile Framework: ", "Flutter Framework version 3.x with Dart SDK version 3.x.")
    add_bullet(doc, "Web Framework: ", "Flutter Web / HTML5 / Modern ECMAScript 6+ standard.")
    add_bullet(doc, "Web Browsers: ", "Google Chrome 90+, Mozilla Firefox 88+, Microsoft Edge 90+, or Apple Safari 14+.")
    
    add_heading_3(doc, "2.2.2 Backend, Database & Cloud Services")
    add_bullet(doc, "Backend Infrastructure: ", "Supabase Cloud PaaS (PostgreSQL 15 engine with PostGIS geospatial extension).")
    add_bullet(doc, "Security & Auth: ", "Supabase Auth with JSON Web Tokens (JWT) and PostgreSQL Row Level Security (RLS) policies.")
    add_bullet(doc, "Cloud Storage: ", "Supabase S3-compatible Object Storage for encrypted KYC documents, spot photos, and verification videos.")
    add_bullet(doc, "Mapping & Geospatial API: ", "Google Maps Platform SDK (Maps JavaScript API, Places API, Geocoding API, Directions API).")
    add_bullet(doc, "Payment Gateway: ", "SSLCommerz Payment Gateway SDK (Sandbox & Live modes supporting MFS and Card networks).")
    add_bullet(doc, "Push Notification Engine: ", "Firebase Cloud Messaging (FCM) & Flutter Local Notifications.")
    add_bullet(doc, "On-Device ML & Scanner: ", "Google ML Kit Text Recognition & Mobile Scanner plugin for instant QR decoding.")

    # =========================================================================
    # 3. INSTALLATION AND SETUP
    # =========================================================================
    add_heading_1(doc, "3. Installation and Setup")
    add_body_p(doc,
        "Setting up ParkFinity is designed to be straightforward for all user categories. Follow the sequential procedures "
        "outlined below to configure the mobile application or administrative portal."
    )
    
    add_heading_2(doc, "3.1 Mobile Application Installation (Rider & Owner)")
    add_bullet(doc, "Step 1 (Obtain APK): ", "Download the official `parkfinity.apk` installation package from the project repository, official website, or distributed release build.")
    add_bullet(doc, "Step 2 (Enable Unknown Sources): ", "If side-loading on Android, navigate to Settings > Security > Install Unknown Apps, and grant installation permission to your browser/file manager.")
    add_bullet(doc, "Step 3 (Execute Installation): ", "Tap on the downloaded `parkfinity.apk` file and tap 'Install'. The package installer will verify the build and complete installation within seconds.")
    add_bullet(doc, "Step 4 (Launch & Grant Permissions): ", "Open the installed ParkFinity application from your app drawer. When prompted, grant the following essential runtime permissions:")
    add_bullet(doc, "  - Location Permission: ", "Select 'While using the app' and enable 'Precise Location' so the app can locate nearby parking spots.")
    add_bullet(doc, "  - Camera Permission: ", "Required for instant QR code scanning at parking entry/exit and capturing KYC verification photographs.")
    add_bullet(doc, "  - Storage/Media Permission: ", "Required to upload spot gallery photos, property ownership documents, and driving license images.")
    add_bullet(doc, "  - Notification Permission: ", "Required to receive critical real-time alerts regarding booking confirmations, check-in reminders, and timer expiration.")

    add_heading_2(doc, "3.2 Web Administrative Portal Access")
    add_bullet(doc, "Step 1 (Open Browser): ", "Launch a modern, standards-compliant web browser (Google Chrome, Mozilla Firefox, or Microsoft Edge).")
    add_bullet(doc, "Step 2 (Navigate to Portal URL): ", "Enter the administrative URL (`https://admin.parkfinity.com` or local deployment address `http://localhost:port`).")
    add_bullet(doc, "Step 3 (Security Check): ", "Ensure the connection shows a valid SSL/TLS padlock icon in the browser address bar to guarantee end-to-end data encryption.")

    add_heading_2(doc, "3.3 Initial System Configuration")
    add_body_p(doc,
        "Upon first opening the mobile application or admin portal, ensure your device has a stable internet connection. "
        "The application automatically fetches initial platform configuration parameters, dynamic pricing rules, active "
        "geographical zones, and server status from Supabase Cloud."
    )

    # =========================================================================
    # 4. GETTING STARTED
    # =========================================================================
    add_heading_1(doc, "4. Getting Started")
    add_body_p(doc,
        "This section provides an introductory overview of the ParkFinity user journey, from launching the application "
        "for the first time to choosing your operational role and navigating the core dashboard."
    )
    
    add_heading_2(doc, "4.1 Launching the Application")
    add_body_p(doc,
        "When you tap the ParkFinity icon on your mobile device, the animated splash screen appears while the app initializes "
        "local state, checks network connectivity, and validates existing authentication tokens."
    )
    
    add_heading_2(doc, "4.2 Role Selection & Onboarding")
    add_body_p(doc,
        "If you are a first-time user, ParkFinity presents the 'Choose Your Path' screen where you select your primary role:"
    )
    add_bullet(doc, "Find Parking (Rider): ", "Select this option if you own a vehicle (car, motorcycle, pickup, SUV) and wish to locate, reserve, and park at verified parking spots across the city.")
    add_bullet(doc, "Host Parking (Space Owner): ", "Select this option if you own or manage residential garages, commercial parking spaces, or vacant land and want to earn money hosting vehicles.")

    add_heading_2(doc, "4.3 Registration and Authentication")
    add_body_p(doc,
        "Creating an account requires your full name, active mobile phone number, and a secure password. Alternatively, "
        "you can utilize one-tap Google Authentication. Once registered, riders and owners complete an identity verification "
        "(KYC) step to ensure trust and security across the platform."
    )
    
    add_heading_2(doc, "4.4 Navigation & Bottom Bar Overview")
    add_body_p(doc,
        "The mobile interface provides a persistent, modern bottom navigation bar tailored to each user role:"
    )
    add_bullet(doc, "Rider Navigation Bar: ", "Consists of Explore (Interactive Map & Search), Bookings (Active sessions & history), Garage (Vehicle management), and Profile (Settings, wallet & KYC).")
    add_bullet(doc, "Owner Navigation Bar: ", "Consists of Dashboard (Revenue & summary), Bookings (Incoming reservations), My Listings (Spot capacity & QR), Wallet (Payouts & balance), and Profile.")

    # =========================================================================
    # 5. USER ROLES AND ACCESS PERMISSIONS
    # =========================================================================
    add_heading_1(doc, "5. User Roles and Access Permissions")
    add_body_p(doc,
        "ParkFinity enforces a strict Role-Based Access Control (RBAC) architecture enforced by PostgreSQL Row-Level "
        "Security (RLS) policies. Each user category has specific permissions and dedicated interface screens."
    )
    
    add_heading_2(doc, "5.1 System Administrator")
    add_body_p(doc,
        "The System Administrator holds superuser governance capabilities across the entire ParkFinity platform via the "
        "dedicated Web Administrative Portal. Key responsibilities and permissions include:"
    )
    add_bullet(doc, "User Account Governance: ", "View complete profiles of all registered riders and hosts, inspect account status, and suspend or reactivate accounts.")
    add_bullet(doc, "KYC Compliance & Verification: ", "Inspect uploaded National Identity Cards, driving licenses, and property ownership deeds; approve or reject verification requests with feedback.")
    add_bullet(doc, "Parking Listings Moderation: ", "Review all published parking listings, verify spot details and videos, and suspend fraudulent or substandard spots.")
    add_bullet(doc, "Live Booking Supervision: ", "Monitor real-time active parking sessions, detect slot overstays, and supervise manual slot releases.")
    add_bullet(doc, "Financial Audit & Payouts: ", "Audit all wallet top-ups, booking fees, platform commission deductions, and process host withdrawal requests.")
    add_bullet(doc, "Dynamic Policy Configuration: ", "Configure platform commission rates (e.g., 10%), peak-hour price multipliers (e.g., 1.5×), and cancellation refund parameters.")

    add_heading_2(doc, "5.2 Parking Space Owner (Host)")
    add_body_p(doc,
        "Parking Space Owners are verified individuals or organizations that list and monetize their parking spots. "
        "Their capabilities include:"
    )
    add_bullet(doc, "Spot Creation & Management: ", "Create detailed parking listings with photos, verification video, GPS pin, vehicle slot capacities, and security amenities.")
    add_bullet(doc, "Pricing & Availability Schedule: ", "Define custom hourly, daily, weekly, and monthly tariffs; configure weekly operating hours per day (e.g., Sat-Thu 08:00–20:00).")
    add_bullet(doc, "Spot QR Code Standee: ", "Generate and download a high-resolution printable QR code and 6-character short code for physical mounting at the garage entrance.")
    add_bullet(doc, "Booking Monitoring: ", "View upcoming, active, and completed driver bookings, check-in timestamps, and vehicle registration numbers.")
    add_bullet(doc, "Earnings & Bank Withdrawal: ", "Track accumulated net income in real time and submit withdrawal requests directly to their bank account.")

    add_heading_2(doc, "5.3 Rider / Vehicle Driver")
    add_body_p(doc,
        "Riders are registered vehicle drivers seeking safe, affordable, and guaranteed parking. Their capabilities include:"
    )
    add_bullet(doc, "Map & List Search: ", "Locate nearby parking spots via interactive Google Maps or searchable list view with real-time slot availability counts.")
    add_bullet(doc, "AI Smart Assistant: ", "Query the AI recommendation assistant in natural language to receive context-ranked recommendations.")
    add_bullet(doc, "Garage & Vehicle Registry: ", "Manage multiple vehicles (cars, motorcycles, pickups) with license plate validation.")
    add_bullet(doc, "Instant Booking & Wallet Payment: ", "Top up digital wallet via SSLCommerz (bKash, Nagad, Cards) and execute one-touch instant parking reservations.")
    add_bullet(doc, "QR Check-in & Live Countdown: ", "Scan the spot's physical QR code upon arrival to start parking, monitor real-time countdown timer, extend time, and scan again to checkout.")
    add_bullet(doc, "Booking History & Reviews: ", "Access complete booking history, review digital receipts, and submit ratings/feedback for hosted parking spaces.")

    add_heading_2(doc, "5.4 Access Control Matrix")
    add_body_p(doc,
        "The following matrix summarizes the comparative operational permissions across all three platform roles:"
    )
    
    # RBAC Table
    headers = ["Feature / Operation", "Rider", "Space Owner", "System Admin"]
    widths = [2.6, 1.3, 1.3, 1.3]
    matrix_data = [
        ["Search Parking Spots on Map", "Yes", "Yes", "Yes (Audit)"],
        ["AI Natural Language Parking Query", "Yes", "No", "No"],
        ["Add & Manage Personal Vehicles", "Yes", "No", "No"],
        ["Book Parking Spot & Pay via Wallet", "Yes", "No", "No"],
        ["Scan Spot QR Code (Check-in/Out)", "Yes", "No", "No"],
        ["List New Parking Space (Photos/Video)", "No", "Yes", "Supervise"],
        ["Generate Spot QR Standee Code", "No", "Yes", "Supervise"],
        ["Configure Pricing & Schedules", "No", "Yes", "Override"],
        ["Request Bank Revenue Withdrawal", "No", "Yes", "Approve/Reject"],
        ["Submit KYC Verification Docs", "Yes (NID/License)", "Yes (NID/Deed)", "Vetting Only"],
        ["Approve / Reject KYC Submissions", "No", "No", "Yes"],
        ["Suspend / Reactivate User Accounts", "No", "No", "Yes"],
        ["Configure Platform Commission & Multipliers", "No", "No", "Yes"],
        ["View Platform Financial Reports", "Personal Wallet", "Owner Wallet", "Full Ledger"]
    ]
    tbl_rbac = doc.add_table(rows=1, cols=4)
    format_styled_table(tbl_rbac, widths, headers, matrix_data)
    
    add_body_p(doc, "") # Spacing

print("Sections 1 to 5 module ready.")
