import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .common import (
    COLOR_PRIMARY, COLOR_SECONDARY, COLOR_DARK, COLOR_MUTED,
    add_heading_1, add_heading_2, add_heading_3,
    add_body_p, add_bullet, add_step, add_fig, add_callout
)

def add_section_6(doc):
    # =========================================================================
    # 6. FUNCTIONALITIES AND USAGE INSTRUCTIONS
    # =========================================================================
    add_heading_1(doc, "6. Functionalities and Usage Instructions")
    add_body_p(doc,
        "This section constitutes the primary operational guide of the ParkFinity User Manual. It provides comprehensive, "
        "step-by-step instructions accompanied by high-resolution visual screenshots for all user operations across the "
        "Onboarding & Authentication, Rider Client, Parking Space Owner Host, and System Administrator modules. "
        "Every figure is paired with an exact descriptive caption and a contextual one-line explanation detailing its "
        "role within the application workflow."
    )
    
    # -------------------------------------------------------------------------
    # 6.1 ONBOARDING & KYC IDENTITY VERIFICATION
    # -------------------------------------------------------------------------
    add_heading_2(doc, "6.1 Onboarding & KYC Identity Verification")
    add_body_p(doc,
        "Before engaging in parking transactions or space hosting, every user undergoes role assignment, registration, "
        "and legal identity verification to maintain platform integrity, mutual trust, and accountability."
    )
    
    # 6.1.1 Role Selection
    add_heading_3(doc, "6.1.1 Role Selection Workflow")
    add_body_p(doc,
        "Upon launching the ParkFinity mobile application for the first time, users are presented with a role selection "
        "screen that determines their user journey and unlocks the respective toolsets."
    )
    add_step(doc, 1, "Select Role", "Tap 'Find Parking' if you are a driver seeking parking spaces, or tap 'Host Parking' if you own an available parking space.")
    add_step(doc, 2, "Proceed to Auth", "The application automatically configures your onboarding profile and navigates to the Registration/Login portal.")
    add_fig(doc, "rider_images/IMG_20260815_145026_169.jpg",
            "Figure 6.1: Role Selection Screen",
            "Allows the new user to choose whether to register as a vehicle driver looking for parking or as a property owner hosting a parking space.",
            is_mobile=True)
    
    # 6.1.2 User Registration
    add_heading_3(doc, "6.1.2 User Registration & Account Creation")
    add_body_p(doc,
        "New users can register a secure personal account using their basic contact details or Google Single Sign-On (SSO)."
    )
    add_step(doc, 1, "Enter Personal Information", "Input your Full Name, valid Bangladeshi Mobile Phone Number (+880), and a secure password (minimum 8 characters).")
    add_step(doc, 2, "Submit Registration", "Tap the 'Sign Up' button to submit your credentials.")
    add_step(doc, 3, "Google Sign-In Alternative", "Alternatively, tap 'Continue with Google' to instantly authenticate using your verified Google account.")
    add_fig(doc, "owner_images/Screenshot_2026-08-15-14-45-57-351_com.parkfinity.parkfinity-edit.jpg",
            "Figure 6.2: User Registration and Account Creation Screen",
            "Enables new users to register their account using their full name, mobile phone number, secure password, or Google authentication.",
            is_mobile=True)
    
    # 6.1.3 Rider KYC Verification
    add_heading_3(doc, "6.1.3 Rider KYC Identity Verification")
    add_body_p(doc,
        "To prevent vehicle theft, unauthorized bookings, and identity fraud, vehicle drivers must submit valid government "
        "identification before booking their first parking spot."
    )
    add_step(doc, 1, "Access KYC Screen", "Navigate to the KYC verification screen after registration or via Profile > Verification Status.")
    add_step(doc, 2, "Upload NID Front & Back", "Tap to capture or upload clear photographs of the front and backside of your National Identity Card.")
    add_step(doc, 3, "Upload Driving License", "Capture and upload a crisp photograph of your official BRTA Driving License.")
    add_step(doc, 4, "Submit for Review", "Tap 'Submit KYC' to securely transmit your encrypted documents to the administrative review queue.")
    add_fig(doc, "rider_images/IMG_20260815_145026_277.jpg",
            "Figure 6.3: Rider KYC Identity Verification Screen",
            "Allows drivers to upload both sides of their National Identity Card (NID) and driving license for identity validation and fraud prevention.",
            is_mobile=True)
    
    # 6.1.4 Owner KYC Verification
    add_heading_3(doc, "6.1.4 Space Owner KYC & Property Ownership Verification")
    add_body_p(doc,
        "Parking space hosts must provide legal proof of property ownership or authorized tenancy alongside personal NID "
        "to ensure listed parking spaces are legitimate and legally compliant."
    )
    add_step(doc, 1, "Initiate Owner KYC", "Open the Host KYC verification screen during space registration.")
    add_step(doc, 2, "Upload Government NID", "Upload clear photographs of both sides of your National Identity Card.")
    add_step(doc, 3, "Upload Property Ownership Document", "Upload a utility bill (electricity/gas), property deed, or building management allotment document proving space rights.")
    add_step(doc, 4, "Submit for Compliance Vetting", "Tap 'Submit' to send documents to the administrative team for verification.")
    add_fig(doc, "owner_images/Screenshot_2026-08-15-14-46-17-316_com.parkfinity.parkfinity-edit.jpg",
            "Figure 6.4: Parking Space Owner KYC Verification Screen",
            "Allows property hosts to upload government NID cards and property ownership documents to legally verify hosting eligibility.",
            is_mobile=True)

    # -------------------------------------------------------------------------
    # 6.2 RIDER MODULE OPERATIONS
    # -------------------------------------------------------------------------
    add_heading_2(doc, "6.2 Rider Module Operations")
    add_body_p(doc,
        "The Rider module provides vehicle drivers with a complete suite of tools to discover, reserve, navigate to, and "
        "pay for parking spaces seamlessly."
    )
    
    # 6.2.1 Interactive Map
    add_heading_3(doc, "6.2.1 Interactive Map & Real-Time Parking Discovery")
    add_body_p(doc,
        "The interactive map screen serves as the primary navigation hub for drivers, displaying live GPS location and "
        "nearby available parking pins."
    )
    add_step(doc, 1, "Explore Nearby Spots", "View color-coded parking pins representing available parking locations in your immediate geographic radius.")
    add_step(doc, 2, "Search Destination", "Tap the top search bar ('Where do you want to park?') and type your destination to inspect parking availability around specific landmarks.")
    add_step(doc, 3, "Inspect Spot Card", "Tap on any map pin to view the parking spot summary card showing spot name, hourly tariff, available slots, and distance.")
    add_fig(doc, "rider_images/IMG_20260815_145022_252.jpg",
            "Figure 6.5: Interactive Map and Nearby Parking Discovery Screen",
            "Displays real-time geospatial map pins, available parking spots, active session shortcuts, and intuitive search destination queries.",
            is_mobile=True)

    # 6.2.2 Parking Spot Listings
    add_heading_3(doc, "6.2.2 Parking Spot Listings & Filter Criteria")
    add_body_p(doc,
        "Drivers can switch to a structured list view to compare multiple parking spots simultaneously."
    )
    add_step(doc, 1, "Open Listings View", "Tap 'All parking spots' from the Explore screen.")
    add_step(doc, 2, "Review Available Spots", "Inspect each parking card displaying the spot title, address, hourly tariff (e.g., ৳100/hr), and free slot counter (e.g., '2 free').")
    add_step(doc, 3, "Filter Criteria", "Tap the filter icon to refine listings by vehicle category (Car, Bike, Pickup), covered roof, CCTV, or maximum hourly price.")
    add_fig(doc, "rider_images/IMG_20260815_145017_594.jpg",
            "Figure 6.6: Parking Spot Listings Screen",
            "Provides a searchable list view of all registered parking spots with hourly rates, distance, and real-time available slot counters.",
            is_mobile=True)

    # 6.2.3 AI Smart Assistant
    add_heading_3(doc, "6.2.3 AI-Powered Smart Parking Natural Language Assistant")
    add_body_p(doc,
        "ParkFinity includes an intelligent AI assistant that allows drivers to express their parking constraints in natural language."
    )
    add_step(doc, 1, "Access AI Assistant", "Tap 'Ask ParkFinity AI' located on the search banner.")
    add_step(doc, 2, "Enter Preferences", "Type your natural-language request (e.g., 'Cheapest covered spot near NSTU campus with CCTV').")
    add_step(doc, 3, "Ask AI", "Tap 'Ask AI' to process your query against real-time listing attributes.")
    add_fig(doc, "rider_images/IMG_20260815_145017_961.jpg",
            "Figure 6.7: AI-Powered Smart Parking Assistant Screen",
            "Allows drivers to enter natural-language parking requirements to receive personalized, context-aware parking suggestions.",
            is_mobile=True)

    # 6.2.4 AI Smart Recommendations
    add_heading_3(doc, "6.2.4 AI Smart Recommendations & Decision Rationales")
    add_body_p(doc,
        "The AI recommendation service ranks candidate parking spots and provides transparent rationales for each recommendation."
    )
    add_step(doc, 1, "Review Top Pick", "Inspect the 'Top Pick' parking spot highlighting proximity, price per hour, and total vacant slots.")
    add_step(doc, 2, "Read AI Reasoning", "Review the AI-generated rationale explaining why this spot is optimal based on distance, safety, and price.")
    add_step(doc, 3, "Proceed to Booking", "Tap the recommended card to immediately proceed to spot checkout.")
    add_fig(doc, "rider_images/IMG_20260815_145018_217.jpg",
            "Figure 6.8: AI Smart Recommendations Result Screen",
            "Presents top-ranked parking spot recommendations along with AI-generated reasoning regarding proximity, pricing, and available amenities.",
            is_mobile=True)

    # 6.2.5 Vehicle Management
    add_heading_3(doc, "6.2.5 Vehicle Management (My Garage)")
    add_body_p(doc,
        "The Garage module allows drivers to maintain a fleet of personal vehicles for quick selection during booking."
    )
    add_step(doc, 1, "Open Garage", "Tap 'Garage' in the bottom navigation bar.")
    add_step(doc, 2, "View Registered Vehicles", "Inspect saved cars and motorcycles with manufacturer brand, model, and official license plate numbers (e.g., Dhaka-Metro-Ga-12-3478).")
    add_step(doc, 3, "Add New Vehicle", "Tap '+ Add Vehicle' to register an additional automobile.")
    add_fig(doc, "rider_images/IMG_20260815_145017_736.jpg",
            "Figure 6.9: Rider Garage (My Vehicles) Management Screen",
            "Displays the user's registered vehicles and license plates with options to switch active vehicles or register additional automobiles.",
            is_mobile=True)

    # 6.2.6 Add New Vehicle
    add_heading_3(doc, "6.2.6 Adding a New Vehicle to Garage")
    add_body_p(doc,
        "Registering a new vehicle takes less than a minute and ensures slot dimensions match your vehicle type."
    )
    add_step(doc, 1, "Select Vehicle Type", "Choose vehicle category: Car, Motorcycle, SUV, or Pickup.")
    add_step(doc, 2, "Enter Vehicle Details", "Enter Brand (e.g., Toyota), Model (e.g., Corolla), and License Plate number.")
    add_step(doc, 3, "Save Vehicle", "Tap 'Save Vehicle' to store the vehicle in your cloud garage profile.")
    add_fig(doc, "rider_images/IMG_20260815_145017_907.jpg",
            "Figure 6.10: Add New Vehicle Screen",
            "Enables drivers to register a new vehicle by specifying the vehicle category, manufacturer brand, model, and official license plate number.",
            is_mobile=True)

    # 6.2.7 Digital Wallet & Top-up
    add_heading_3(doc, "6.2.7 Digital Wallet & Top-up via SSLCommerz")
    add_body_p(doc,
        "ParkFinity includes an integrated digital wallet that eliminates cash handling and ensures instant automated parking deductions."
    )
    add_step(doc, 1, "Check Balance", "View current available balance (e.g., ৳21,920.00) and recent transaction history.")
    add_step(doc, 2, "Add Funds", "Tap '+ Add Funds', enter amount, and select payment method via SSLCommerz gateway (bKash, Nagad, Rocket, Visa, Mastercard).")
    add_step(doc, 3, "Audit Ledger", "Review completed booking deductions, top-up credits, and automatic cancellation refunds.")
    add_fig(doc, "rider_images/IMG_20260815_145017_909.jpg",
            "Figure 6.11: Digital Wallet and Top-up Screen",
            "Displays available wallet balance, quick funds addition via SSLCommerz gateway, and an itemized ledger of booking deductions and refunds.",
            is_mobile=True)

    # 6.2.8 Active Parking Session
    add_heading_3(doc, "6.2.8 Active Parking Session & Live Countdown Timer")
    add_body_p(doc,
        "Once a driver arrives at the parking spot and scans the check-in QR code, the active session screen activates."
    )
    add_step(doc, 1, "Monitor Live Timer", "Track remaining parking duration in real time (e.g., 'Time Remaining: 5:48').")
    add_step(doc, 2, "Extend Session", "Tap 'Extend' to add extra parking hours and pay automatically from wallet balance.")
    add_step(doc, 3, "Navigate to Spot", "Tap 'Navigate' to launch turn-by-turn Google Maps navigation directly to the parking gate.")
    add_step(doc, 4, "Check out", "Tap 'Check out' and scan the spot's QR code to conclude the session and release the parking slot.")
    add_fig(doc, "rider_images/IMG_20260815_145018_100.jpg",
            "Figure 6.12: Active Parking Session and Live Countdown Timer Screen",
            "Provides real-time countdown of remaining parking time with one-touch options to extend duration, navigate to spot, or checkout via QR code.",
            is_mobile=True)

    # 6.2.9 Booking History
    add_heading_3(doc, "6.2.9 Rider Booking History & Receipts")
    add_body_p(doc,
        "The Bookings tab maintains an exhaustive archive of all past, active, and upcoming parking reservations."
    )
    add_step(doc, 1, "Access Bookings", "Tap 'Bookings' from the bottom navigation bar.")
    add_step(doc, 2, "Inspect Sessions", "View date, time intervals, spot location names (e.g., 'ASH - NSTU'), and status tags.")
    add_step(doc, 3, "View Receipt", "Tap any booking record to inspect complete cost breakdown and host contact details.")
    add_fig(doc, "rider_images/IMG_20260815_145018_405.jpg",
            "Figure 6.13: Rider Booking History and Records Screen",
            "Maintains a chronological record of all upcoming, active, and past parking sessions with location names, dates, and time intervals.",
            is_mobile=True)

    # 6.2.10 Push Notifications
    add_heading_3(doc, "6.2.10 Rider Push Notifications & Event Alerts")
    add_body_p(doc,
        "Push notifications keep drivers updated regarding vital parking lifecycle events."
    )
    add_step(doc, 1, "View Notifications", "Tap the bell icon in the top app bar to view notification history.")
    add_step(doc, 2, "Review Alerts", "Inspect booking confirmations, check-in timestamps, 15-minute expiration alerts, and cancellation notices.")
    add_step(doc, 3, "Mark as Read", "Tap 'Mark all as read' to clear unread notification badges.")
    add_fig(doc, "rider_images/IMG_20260815_145017_925.jpg",
            "Figure 6.14: Rider Push Notifications Screen",
            "Alerts drivers in real time regarding booking confirmations, session check-in timestamps, impending expirations, and cancellations.",
            is_mobile=True)

    # 6.2.11 Rider Profile
    add_heading_3(doc, "6.2.11 Rider Profile & Account Settings")
    add_body_p(doc,
        "The Profile screen allows drivers to manage their personal details, security settings, and verified identity status."
    )
    add_step(doc, 1, "View Profile Information", "Inspect your name, registered email address, and mobile phone number.")
    add_step(doc, 2, "Check KYC Status", "Verify your identity badge ('Your NID and licence were submitted at sign-up').")
    add_step(doc, 3, "Quick Shortcuts", "Access Edit Profile, My Vehicles, Booking History, and Account Settings.")
    add_fig(doc, "rider_images/IMG_20260815_145017_535.jpg",
            "Figure 6.15: Rider Profile and Account Settings Screen",
            "Displays user profile details, contact information, verification badge, shortcuts to garage, booking history, and logout options.",
            is_mobile=True)

    # -------------------------------------------------------------------------
    # 6.3 PARKING SPACE OWNER MODULE OPERATIONS
    # -------------------------------------------------------------------------
    add_heading_2(doc, "6.3 Parking Space Owner Module Operations")
    add_body_p(doc,
        "The Parking Space Owner (Host) module provides property hosts with comprehensive tools to publish parking spaces, "
        "manage vehicle capacities, track live driver reservations, and withdraw accumulated earnings."
    )
    
    # 6.3.1 Owner Dashboard
    add_heading_3(doc, "6.3.1 Space Owner Dashboard & Earnings Summary")
    add_body_p(doc,
        "The Owner Dashboard serves as the command center for hosts, summarizing daily performance and active parkings."
    )
    add_step(doc, 1, "Monitor Financial Performance", "View Today's Earnings and Total Bookings Ever at a glance.")
    add_step(doc, 2, "Track Active Parkings", "Monitor real-time count of vehicles currently parked in your facilities.")
    add_step(doc, 3, "Quick Actions", "Tap '+ Add New Parking Spot' to register a new space or 'Withdraw Earnings' to initiate bank payout.")
    add_fig(doc, "owner_images/Screenshot_2026-08-15-14-43-29-335_com.parkfinity.parkfinity-edit.jpg",
            "Figure 6.16: Space Owner Dashboard and Earnings Summary Screen",
            "Summarizes today's revenue, all-time bookings, active parked vehicles, and quick action shortcuts for parking spot management.",
            is_mobile=True)

    # 6.3.2 Add Spot - Photos & Video
    add_heading_3(doc, "6.3.2 Adding a Parking Spot – Photos & Verification Video")
    add_body_p(doc,
        "To maintain high listing quality and authenticity, hosts must upload visual proof of their parking facility."
    )
    add_step(doc, 1, "Add Spot Photos", "Upload a minimum of 3 high-resolution photos showing entry gate, parking surface, and surroundings.")
    add_step(doc, 2, "Add Spot Video", "Record and upload a mandatory 15-second walkthrough video demonstrating vehicle entry and clearance.")
    add_step(doc, 3, "Enter Spot Details", "Provide Listing Title (e.g., 'NSTU Central Garage'), Description, and Full Physical Address.")
    add_fig(doc, "owner_images/Screenshot_2026-08-15-14-44-09-446_com.parkfinity.parkfinity-edit.jpg",
            "Figure 6.17: Add Parking Spot – Photos and Video Upload Screen",
            "Guides owners to upload high-quality photographs (minimum 3) and a mandatory verification video of the parking facility.",
            is_mobile=True)

    # 6.3.3 Add Spot - Vehicle Slots, Pricing & Amenities
    add_heading_3(doc, "6.3.3 Adding a Parking Spot – Vehicle Slots, Pricing & Amenities")
    add_body_p(doc,
        "Hosts define slot capacity per vehicle type and set multi-tier pricing structures."
    )
    add_step(doc, 1, "Configure Vehicle Slots", "Define slot counts for Car (e.g., 2 slots), Motorcycle (e.g., 5 slots), and Pickup.")
    add_step(doc, 2, "Set Pricing Options", "Input base tariffs: Hourly (৳), Daily (৳), Weekly (৳), and Monthly (৳).")
    add_step(doc, 3, "Select Security Amenities", "Check available amenities: CCTV Camera, Covered Parking, Security Guard, and EV Charging.")
    add_fig(doc, "owner_images/Screenshot_2026-08-15-14-44-22-153_com.parkfinity.parkfinity-edit.jpg",
            "Figure 6.18: Add Parking Spot – Vehicle Slots, Pricing, and Amenities Screen",
            "Configures slot capacity by vehicle type, multi-tiered pricing (hourly/daily/monthly), and security amenities such as CCTV and EV charging.",
            is_mobile=True)

    # 6.3.4 Add Spot - Booking Mode & Availability Schedule
    add_heading_3(doc, "6.3.4 Adding a Parking Spot – Booking Mode & Availability Schedule")
    add_body_p(doc,
        "Hosts can specify whether reservations require manual host approval and configure daily operating hours."
    )
    add_step(doc, 1, "Choose Booking Mode", "Select 'Instant' (drivers can book immediately) or 'Manual Approval' (host confirms each booking).")
    add_step(doc, 2, "Configure Daily Hours", "Set operational start and end times for each day of the week (Sat through Fri, e.g., 00:00 to 23:59 or customized business hours).")
    add_fig(doc, "owner_images/Screenshot_2026-08-15-14-44-35-252_com.parkfinity.parkfinity-edit.jpg",
            "Figure 6.19: Add Parking Spot – Booking Mode and Availability Schedule Screen",
            "Allows hosts to select instant booking or manual approval mode and define operational time windows across all seven days of the week.",
            is_mobile=True)

    # 6.3.5 Add Spot - Map Location Pinning
    add_heading_3(doc, "6.3.5 Adding a Parking Spot – Map Location Pinning & Publishing")
    add_body_p(doc,
        "Accurate geospatial pinning ensures drivers are routed directly to the exact parking entrance."
    )
    add_step(doc, 1, "Pin Exact Location", "Drag the map pin or tap 'Use my location' to set precise GPS coordinates on Google Maps.")
    add_step(doc, 2, "Verify Placement", "Confirm the address matches the physical entrance gate.")
    add_step(doc, 3, "Publish Listing", "Tap 'Publish Listing' to publish the parking spot to the live ParkFinity platform.")
    add_fig(doc, "owner_images/Screenshot_2026-08-15-14-44-54-387_com.parkfinity.parkfinity-edit.jpg",
            "Figure 6.20: Add Parking Spot – Map Location Pinning and Publishing Screen",
            "Enables precise GPS coordinate placement on Google Maps and final publication of the parking listing to the live platform.",
            is_mobile=True)

    # 6.3.6 My Listings
    add_heading_3(doc, "6.3.6 Managing Parking Listings & Slot Availability")
    add_body_p(doc,
        "The My Listings tab provides hosts with real-time visibility into all their published parking facilities."
    )
    add_step(doc, 1, "View Hosted Spots", "Inspect spot title, full address, and status badge ('Active').")
    add_step(doc, 2, "Monitor Slot Occupancy", "Review live slot utilization per vehicle type (e.g., 'Car: 2/2', 'Pickup: 1/1', 'Motorcycle: 0/1').")
    add_step(doc, 3, "Manage Listing", "Tap on a spot card to edit pricing, pause hosting, or view its dedicated QR code.")
    add_fig(doc, "owner_images/Screenshot_2026-08-15-14-43-56-586_com.parkfinity.parkfinity-edit.jpg",
            "Figure 6.21: My Listings and Slot Capacity Screen",
            "Displays all published parking spots owned by the host with real-time occupancy counts for cars, pickups, and motorcycles.",
            is_mobile=True)

    # 6.3.7 Spot QR Code
    add_heading_3(doc, "6.3.7 Spot QR Code Generation & Physical Standee")
    add_body_p(doc,
        "Every published parking spot is automatically assigned a unique cryptographic QR code and 6-character short code."
    )
    add_step(doc, 1, "Access Spot QR", "Open the spot details and tap 'Spot QR Code'.")
    add_step(doc, 2, "Download & Print", "Tap 'Download QR code' to save a high-resolution printable PDF/image standee.")
    add_step(doc, 3, "Mount at Entrance", "Mount the printed standee at the parking entrance so drivers can scan upon arrival and departure.")
    add_fig(doc, "owner_images/Screenshot_2026-08-15-14-45-38-147_com.parkfinity.parkfinity-edit.jpg",
            "Figure 6.22: Parking Spot QR Code and Check-in Standee Screen",
            "Provides a downloadable and printable high-resolution QR code and alphanumeric short code for physical mounting at the parking entrance.",
            is_mobile=True)

    # 6.3.8 Owner Booking Management
    add_heading_3(doc, "6.3.8 Owner Booking Management & Driver Tracking")
    add_body_p(doc,
        "Hosts can monitor incoming reservations and track parked vehicles in real time."
    )
    add_step(doc, 1, "Open Bookings Tab", "Tap 'Bookings' from the host bottom navigation bar.")
    add_step(doc, 2, "Inspect Driver Sessions", "View driver names (e.g., 'Majedur Rahman'), spot names, dates, and scheduled time intervals.")
    add_step(doc, 3, "Contact Driver", "Tap a booking card to access the driver's phone number if urgent coordination is required.")
    add_fig(doc, "owner_images/Screenshot_2026-08-15-14-43-49-505_com.parkfinity.parkfinity-edit.jpg",
            "Figure 6.23: Owner Booking Management and Driver Tracking Screen",
            "Lists active and upcoming driver reservations with driver names, reserved spot details, and designated time schedules.",
            is_mobile=True)

    # 6.3.9 Owner Wallet & Withdrawal
    add_heading_3(doc, "6.3.9 Owner Wallet & Bank Withdrawal")
    add_body_p(doc,
        "Host booking earnings are automatically credited to the owner wallet following successful driver checkout."
    )
    add_step(doc, 1, "Review Available Balance", "Check your current available balance and all-time total earnings.")
    add_step(doc, 2, "Request Bank Withdrawal", "Tap 'Withdraw', enter desired amount and bank account details (Bank name, Branch, Account number, Routing number).")
    add_step(doc, 3, "Transaction History", "Audit all past earnings credits, commission deductions, and bank payouts.")
    add_fig(doc, "owner_images/Screenshot_2026-08-15-14-45-06-433_com.parkfinity.parkfinity-edit.jpg",
            "Figure 6.24: Owner Wallet and Bank Withdrawal Screen",
            "Shows accumulated net revenue after platform commission and provides payout requests directly to the host's designated bank account.",
            is_mobile=True)

    # 6.3.10 Owner Notifications
    add_heading_3(doc, "6.3.10 Owner Notifications & Overstay Alerts")
    add_body_p(doc,
        "The notification center delivers immediate push alerts regarding garage activity."
    )
    add_step(doc, 1, "Check Alerts", "View notifications for 'New booking', 'Driver check-in', and 'Rider overstayed'.")
    add_step(doc, 2, "Overstay Actions", "When a driver overstays past their booked duration, the system automatically alerts the host and applies penalty tariffs.")
    add_fig(doc, "owner_images/Screenshot_2026-08-15-14-45-26-178_com.parkfinity.parkfinity-edit.jpg",
            "Figure 6.25: Owner Notifications and Overstay Alerts Screen",
            "Delivers real-time notifications for incoming driver bookings, arrival check-ins, departures, and vehicle overstay warnings.",
            is_mobile=True)

    # 6.3.11 Owner Profile
    add_heading_3(doc, "6.3.11 Owner Profile & KYC Status")
    add_body_p(doc,
        "Hosts can update their contact details, review verification status, and manage security options."
    )
    add_step(doc, 1, "Inspect Profile Details", "View host name (e.g., 'Doha'), registered email, phone number (+8801845503086), and 'Owner' role badge.")
    add_step(doc, 2, "Verify KYC Status", "Ensure 'Verification status' shows approved verification for NID and property deed.")
    add_step(doc, 3, "Manage Account", "Access Edit Profile, Booking History, and Settings.")
    add_fig(doc, "owner_images/Screenshot_2026-08-15-14-45-15-335_com.parkfinity.parkfinity-edit.jpg",
            "Figure 6.26: Owner Profile and Verification Status Screen",
            "Displays host credentials, verified KYC badge, contact details, profile modification tools, and account preferences.",
            is_mobile=True)

    # -------------------------------------------------------------------------
    # 6.4 SYSTEM ADMINISTRATOR MODULE OPERATIONS
    # -------------------------------------------------------------------------
    add_heading_2(doc, "6.4 System Administrator Module Operations")
    add_body_p(doc,
        "The System Administrator Web Portal equips platform operators with real-time tools to govern users, moderate listings, "
        "verify legal documents, track payments, analyze platform revenue, and configure dynamic pricing algorithms."
    )
    
    # 6.4.1 Admin Authentication
    add_heading_3(doc, "6.4.1 Admin Portal Authentication")
    add_body_p(doc,
        "The administrative web portal is protected by role-based authentication requiring authorized administrator credentials."
    )
    add_step(doc, 1, "Navigate to Admin Portal", "Open the admin portal in your browser.")
    add_step(doc, 2, "Enter Admin Credentials", "Input your registered administrative email address and password.")
    add_step(doc, 3, "Sign In", "Tap 'Sign In' to authenticate and access the executive dashboard.")
    add_fig(doc, "admin_images/Screenshot 2026-08-15 145738.png",
            "Figure 6.27: Admin Portal Authentication Screen",
            "Secures administrative access through credential verification with role-based security enforcement.",
            is_mobile=False)

    # 6.4.2 Admin Dashboard
    add_heading_3(doc, "6.4.2 Admin Executive Dashboard & Financial Overview")
    add_body_p(doc,
        "The executive dashboard delivers platform-wide key performance indicators (KPIs) in real time."
    )
    add_step(doc, 1, "Financial Metrics", "Inspect Platform Revenue, Gross Volume (e.g., ৳5,1740.00), Owner Payouts, and Refunded totals.")
    add_step(doc, 2, "Platform User Breakdown", "Monitor Total Users count, active Riders count (8), and registered Space Owners count (8).")
    add_step(doc, 3, "Operational Activity", "Review Total Listings (3), Active Listings (3), and Total Completed Bookings (19).")
    add_fig(doc, "admin_images/Screenshot 2026-08-15 145756.png",
            "Figure 6.28: Admin Executive Dashboard and Financial Overview Screen",
            "Provides platform-wide metrics including total revenue, gross volume, active listings, registered users, and total bookings.",
            is_mobile=False)

    # 6.4.3 Users Management
    add_heading_3(doc, "6.4.3 User Account Management & Moderation")
    add_body_p(doc,
        "Administrators can inspect all registered user accounts and enforce safety policies."
    )
    add_step(doc, 1, "Search & Filter Users", "Use the search bar ('Search by name or email') to locate specific user profiles.")
    add_step(doc, 2, "Inspect Account Attributes", "View user Full Name, registered email address, assigned role (Rider / Owner / Admin), and creation date.")
    add_step(doc, 3, "Account Moderation", "Suspend compromised accounts or reset security credentials when required.")
    add_fig(doc, "admin_images/Screenshot 2026-08-15 145816.png",
            "Figure 6.29: User Account Management and Moderation Screen",
            "Allows administrators to search, inspect, moderate, and suspend or reactivate driver and host accounts.",
            is_mobile=False)

    # 6.4.4 KYC Review
    add_heading_3(doc, "6.4.4 KYC Document Review & Verification")
    add_body_p(doc,
        "The KYC Review portal enables compliance officers to audit submitted identification documents."
    )
    add_step(doc, 1, "Filter KYC Queue", "Filter submissions by status: 'All', 'Pending', or 'Verified'.")
    add_step(doc, 2, "Inspect Submitted Documents", "Click on any submission (e.g., 'Md Majed Alvi', 'Doha') to inspect high-resolution NID cards, driving licenses, and property deeds.")
    add_step(doc, 3, "Approve or Reject", "Click 'Approve' to grant verified status or 'Reject' with an explanatory remark.")
    add_fig(doc, "admin_images/Screenshot 2026-08-15 145834.png",
            "Figure 6.30: KYC Document Review and Verification Screen",
            "Enables compliance officers to inspect submitted NID cards, driving licenses, and property deeds to approve or reject KYC requests.",
            is_mobile=False)

    # 6.4.5 Listings Management
    add_heading_3(doc, "6.4.5 Parking Listings Supervision & Moderation")
    add_body_p(doc,
        "Administrators monitor all published parking spaces across geographic zones."
    )
    add_step(doc, 1, "Review Listings Table", "Inspect listing Title (e.g., 'ASH', 'nstu'), full address, hourly rates (e.g., ৳100, ৳150), and operational status.")
    add_step(doc, 2, "Audit Quality", "Inspect spot photos, verification videos, and slot dimensions.")
    add_step(doc, 3, "Take Moderation Action", "Click 'Suspend' to temporarily disable listings that violate community guidelines or fail safety checks.")
    add_fig(doc, "admin_images/Screenshot 2026-08-15 145843.png",
            "Figure 6.31: Parking Listings Supervision and Moderation Screen",
            "Displays all published parking spots with status tags, hourly tariffs, and administrative controls to suspend non-compliant listings.",
            is_mobile=False)

    # 6.4.6 Booking Monitoring
    add_heading_3(doc, "6.4.6 Live Booking Monitoring & Overstay Supervision")
    add_body_p(doc,
        "The real-time booking monitor tracks ongoing parking reservations across the city."
    )
    add_step(doc, 1, "Track Active Bookings", "Inspect live booking records displaying driver name, host spot, and scheduled start/end timestamps.")
    add_step(doc, 2, "Detect Overstaying", "Identify vehicle sessions that have exceeded their booked end time.")
    add_step(doc, 3, "Manual Slot Release", "Intervene to release locked slots if a driver fails to scan the checkout QR code.")
    add_fig(doc, "admin_images/Screenshot 2026-08-15 145854.png",
            "Figure 6.32: Live Booking Monitoring and Overstay Supervision Screen",
            "Tracks active parking sessions, scheduled intervals, driver identities, and overstay violations in real time.",
            is_mobile=False)

    # 6.4.7 Payment Monitoring
    add_heading_3(doc, "6.4.7 Payment Monitoring & Financial Ledger")
    add_body_p(doc,
        "The payment monitoring module provides a complete, auditable ledger of all platform financial movements."
    )
    add_step(doc, 1, "Filter Transactions", "Filter financial logs by transaction type (All, Booking Payment, Wallet Top-up, Owner Payout).")
    add_step(doc, 2, "Inspect Transaction Records", "View User Name (e.g., 'Majedur Rahman'), transaction amount (e.g., ৳120), payment channel, and gateway status.")
    add_step(doc, 3, "Reconcile Gateway", "Cross-verify SSLCommerz transaction references with database ledger entries.")
    add_fig(doc, "admin_images/Screenshot 2026-08-15 145904.png",
            "Figure 6.33: Payment Monitoring and Financial Ledger Screen",
            "Maintains an auditable log of all financial transactions, digital wallet top-ups, booking debits, and platform commissions.",
            is_mobile=False)

    # 6.4.8 Revenue Reports
    add_heading_3(doc, "6.4.8 Revenue Reports & Analytics")
    add_body_p(doc,
        "The reporting portal generates periodic financial charts and commission analytics."
    )
    add_step(doc, 1, "Select Timeframe", "Toggle reporting range between '30d' and '90d'.")
    add_step(doc, 2, "Review Revenue Summary", "Inspect Commission Revenue, Gross Transaction Volume, Owner Payouts, and Total Bookings.")
    add_step(doc, 3, "Analyze Daily Trends", "Review the Daily Commission chart to identify peak usage days and revenue patterns.")
    add_fig(doc, "admin_images/Screenshot 2026-08-15 145918.png",
            "Figure 6.34: Revenue Reports and Analytics Screen",
            "Generates periodic financial analytics, commission summaries, and visual revenue trends over 30-day and 90-day timeframes.",
            is_mobile=False)

    # 6.4.9 Platform Settings
    add_heading_3(doc, "6.4.9 Platform Configuration & Dynamic Pricing Policies")
    add_body_p(doc,
        "Administrators can adjust dynamic server-side pricing parameters without requiring mobile app updates."
    )
    add_step(doc, 1, "Configure Commission Rate", "Set the platform commission rate (e.g., `0.1` representing a 10% platform fee per booking).")
    add_step(doc, 2, "Configure Peak Multiplier", "Set peak-hour price multipliers (e.g., `1.5` applied automatically during heavy congestion windows).")
    add_step(doc, 3, "Configure Refund Rules", "Define cancellation grace periods (e.g., 100% refund if cancelled >30 mins before start).")
    add_step(doc, 4, "Save Changes", "Save changes to immediately propagate dynamic pricing across the platform.")
    add_fig(doc, "admin_images/Screenshot 2026-08-15 145939.png",
            "Figure 6.35: Platform Configuration and Dynamic Pricing Policies Screen",
            "Allows administrators to configure platform commission rates, peak-hour multipliers, and automated refund parameters.",
            is_mobile=False)

print("Section 6 module ready.")
