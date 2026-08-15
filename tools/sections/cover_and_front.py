import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from .common import (
    COLOR_PRIMARY, COLOR_SECONDARY, COLOR_DARK, COLOR_MUTED,
    set_cell_margins, set_cell_background, set_table_borders,
    add_heading_1, add_heading_2, add_body_p
)

def add_cover_page(doc, cover_img_path=r"D:\SPL_2\cover_image.png"):
    # Top spacing
    p_top = doc.add_paragraph()
    p_top.paragraph_format.space_before = Pt(6)
    p_top.paragraph_format.space_after = Pt(2)
    p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r_inst = p_top.add_run("INSTITUTE OF INFORMATION TECHNOLOGY\nNOAKHALI SCIENCE AND TECHNOLOGY UNIVERSITY")
    r_inst.bold = True
    r_inst.font.name = "Times New Roman"
    r_inst.font.size = Pt(13.5)
    r_inst.font.color.rgb = COLOR_PRIMARY
    
    p_deg = doc.add_paragraph()
    p_deg.paragraph_format.space_before = Pt(3)
    p_deg.paragraph_format.space_after = Pt(12)
    p_deg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_deg = p_deg.add_run("Bachelor of Science in Software Engineering\nCourse: Software Project Lab-II (Course Code: SE 3112)")
    r_deg.font.name = "Times New Roman"
    r_deg.font.size = Pt(10.5)
    r_deg.font.italic = True
    r_deg.font.color.rgb = COLOR_MUTED
    
    # Cover Logo / Image
    if cover_img_path and os.path.exists(cover_img_path):
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_before = Pt(2)
        p_logo.paragraph_format.space_after = Pt(14)
        p_logo.paragraph_format.keep_with_next = True
        run_logo = p_logo.add_run()
        run_logo.add_picture(cover_img_path, width=Inches(1.8))

    # Title Box
    tbl_title = doc.add_table(rows=1, cols=1)
    tbl_title.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl_title.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_margins(cell, top=200, bottom=200, left=200, right=200)
    set_cell_background(cell, "F0F4F8")
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="16" w:space="0" w:color="1E3A8A"/>'
        f'  <w:bottom w:val="single" w:sz="16" w:space="0" w:color="1E3A8A"/>'
        f'  <w:left w:val="single" w:sz="16" w:space="0" w:color="1E3A8A"/>'
        f'  <w:right w:val="single" w:sz="16" w:space="0" w:color="1E3A8A"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p_t = cell.paragraphs[0]
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t.paragraph_format.space_before = Pt(4)
    p_t.paragraph_format.space_after = Pt(6)
    r_title = p_t.add_run("ParkFinity\nA Smart Platform for Parking Space Optimization")
    r_title.bold = True
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(20)
    r_title.font.color.rgb = COLOR_PRIMARY
    
    p_sub = cell.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(4)
    p_sub.paragraph_format.space_after = Pt(4)
    r_sub = p_sub.add_run("SPL-2 FINAL PROJECT USER MANUAL")
    r_sub.bold = True
    r_sub.font.name = "Times New Roman"
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = COLOR_SECONDARY
    
    # Spacing between title and metadata
    p_mid = doc.add_paragraph()
    p_mid.paragraph_format.space_before = Pt(36)
    p_mid.paragraph_format.space_after = Pt(12)
    
    # Metadata Table (Submitted By & Supervised By)
    tbl_meta = doc.add_table(rows=1, cols=2)
    tbl_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_sub = tbl_meta.cell(0, 0)
    cell_sup = tbl_meta.cell(0, 1)
    cell_sub.width = Inches(3.25)
    cell_sup.width = Inches(3.25)
    set_cell_margins(cell_sub, top=140, bottom=140, left=140, right=140)
    set_cell_margins(cell_sup, top=140, bottom=140, left=140, right=140)
    set_cell_background(cell_sub, "FAFAFA")
    set_cell_background(cell_sup, "FAFAFA")
    
    for c in [cell_sub, cell_sup]:
        tcPr = c._tc.get_or_add_tcPr()
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="6" w:space="0" w:color="D3D3D3"/>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="0" w:color="D3D3D3"/>'
            f'  <w:left w:val="single" w:sz="6" w:space="0" w:color="D3D3D3"/>'
            f'  <w:right w:val="single" w:sz="6" w:space="0" w:color="D3D3D3"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
        
    p_s1 = cell_sub.paragraphs[0]
    p_s1.paragraph_format.space_before = Pt(0)
    p_s1.paragraph_format.space_after = Pt(4)
    r_hdr1 = p_s1.add_run("Submitted By:")
    r_hdr1.bold = True
    r_hdr1.font.name = "Times New Roman"
    r_hdr1.font.size = Pt(11)
    r_hdr1.font.color.rgb = COLOR_PRIMARY
    
    students = [
        ("Md. Shamsuddoha", "ASH2325005M"),
        ("Eftekar Hossen", "ASH2325006M"),
        ("Md. Majedur Rahman", "ASH2125021M")
    ]
    for name, sid in students:
        p_st = cell_sub.add_paragraph()
        p_st.paragraph_format.space_before = Pt(0)
        p_st.paragraph_format.space_after = Pt(2)
        p_st.paragraph_format.line_spacing = 1.15
        r_n = p_st.add_run(f"• {name}\n")
        r_n.bold = True
        r_n.font.name = "Times New Roman"
        r_n.font.size = Pt(10)
        r_i = p_st.add_run(f"   ID: {sid}")
        r_i.font.name = "Times New Roman"
        r_i.font.size = Pt(9.5)
        r_i.font.color.rgb = COLOR_MUTED
        
    p_s2 = cell_sup.paragraphs[0]
    p_s2.paragraph_format.space_before = Pt(0)
    p_s2.paragraph_format.space_after = Pt(4)
    r_hdr2 = p_s2.add_run("Supervised By:")
    r_hdr2.bold = True
    r_hdr2.font.name = "Times New Roman"
    r_hdr2.font.size = Pt(11)
    r_hdr2.font.color.rgb = COLOR_PRIMARY
    
    p_sup_txt = cell_sup.add_paragraph()
    p_sup_txt.paragraph_format.space_before = Pt(0)
    p_sup_txt.paragraph_format.space_after = Pt(2)
    p_sup_txt.paragraph_format.line_spacing = 1.15
    r_sname = p_sup_txt.add_run("Md Hasan Imam\n")
    r_sname.bold = True
    r_sname.font.name = "Times New Roman"
    r_sname.font.size = Pt(10)
    r_sdes = p_sup_txt.add_run("Assistant Professor\nInstitute of Information Technology (IIT)\nNoakhali Science and Technology University (NSTU)")
    r_sdes.font.name = "Times New Roman"
    r_sdes.font.size = Pt(9.5)
    r_sdes.font.color.rgb = COLOR_MUTED
    
    # Bottom Date
    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_before = Pt(48)
    p_date.paragraph_format.space_after = Pt(0)
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_dt = p_date.add_run("Submission Date: August 2026")
    r_dt.bold = True
    r_dt.font.name = "Times New Roman"
    r_dt.font.size = Pt(11)
    r_dt.font.color.rgb = COLOR_DARK
    
    doc.add_page_break()

def add_table_of_contents(doc, toc_entries):
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(12)
    p_t.paragraph_format.space_after = Pt(12)
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run("TABLE OF CONTENTS")
    r_t.bold = True
    r_t.font.name = "Times New Roman"
    r_t.font.size = Pt(16)
    r_t.font.color.rgb = COLOR_PRIMARY
    
    # Structured Clean Table for TOC
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="E0E0E0", sz="2", val="single")
    
    for title, page, level in toc_entries:
        row = tbl.add_row()
        c_title, c_page = row.cells[0], row.cells[1]
        c_title.width = Inches(5.7)
        c_page.width = Inches(0.8)
        set_cell_margins(c_title, top=40, bottom=40, left=60, right=60)
        set_cell_margins(c_page, top=40, bottom=40, left=60, right=60)
        
        p1 = c_title.paragraphs[0]
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after = Pt(1)
        p1.paragraph_format.line_spacing = 1.15
        
        if level == 1:
            p1.paragraph_format.left_indent = Inches(0.0)
            r1 = p1.add_run(title)
            r1.bold = True
            r1.font.name = "Times New Roman"
            r1.font.size = Pt(10.5)
            r1.font.color.rgb = COLOR_PRIMARY
        elif level == 2:
            p1.paragraph_format.left_indent = Inches(0.25)
            r1 = p1.add_run(title)
            r1.font.name = "Times New Roman"
            r1.font.size = Pt(10)
            r1.font.color.rgb = COLOR_DARK
        else:
            p1.paragraph_format.left_indent = Inches(0.45)
            r1 = p1.add_run(title)
            r1.font.name = "Times New Roman"
            r1.font.size = Pt(9.5)
            r1.font.color.rgb = COLOR_MUTED
            
        p2 = c_page.paragraphs[0]
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(1)
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run(str(page))
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(10)
        if level == 1:
            r2.bold = True
            r2.font.color.rgb = COLOR_PRIMARY
        else:
            r2.font.color.rgb = COLOR_DARK

    doc.add_page_break()

def add_list_of_figures(doc, figure_entries):
    p_t = doc.add_paragraph()
    p_t.paragraph_format.space_before = Pt(12)
    p_t.paragraph_format.space_after = Pt(12)
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run("LIST OF FIGURES")
    r_t.bold = True
    r_t.font.name = "Times New Roman"
    r_t.font.size = Pt(16)
    r_t.font.color.rgb = COLOR_PRIMARY
    
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tbl, color="E0E0E0", sz="2", val="single")
    
    for fnum, fcaption, fpage in figure_entries:
        row = tbl.add_row()
        c_title, c_page = row.cells[0], row.cells[1]
        c_title.width = Inches(5.7)
        c_page.width = Inches(0.8)
        set_cell_margins(c_title, top=35, bottom=35, left=60, right=60)
        set_cell_margins(c_page, top=35, bottom=35, left=60, right=60)
        
        p1 = c_title.paragraphs[0]
        p1.paragraph_format.space_before = Pt(1)
        p1.paragraph_format.space_after = Pt(1)
        p1.paragraph_format.line_spacing = 1.15
        
        r_num = p1.add_run(f"Figure {fnum}: ")
        r_num.bold = True
        r_num.font.name = "Times New Roman"
        r_num.font.size = Pt(9.5)
        r_num.font.color.rgb = COLOR_PRIMARY
        
        r_cap = p1.add_run(fcaption)
        r_cap.font.name = "Times New Roman"
        r_cap.font.size = Pt(9.5)
        r_cap.font.color.rgb = COLOR_DARK
        
        p2 = c_page.paragraphs[0]
        p2.paragraph_format.space_before = Pt(1)
        p2.paragraph_format.space_after = Pt(1)
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r2 = p2.add_run(str(fpage))
        r2.font.name = "Times New Roman"
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = COLOR_DARK

    doc.add_page_break()
