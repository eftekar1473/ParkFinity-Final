from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from .common import (
    COLOR_PRIMARY, COLOR_SECONDARY, COLOR_DARK, COLOR_MUTED,
    add_heading_1, add_heading_2, add_heading_3,
    add_body_p, add_bullet, add_step, add_callout, format_styled_table,
    set_cell_margins, set_cell_background, set_table_borders
)

def add_sections_7_to_12(doc):
    # =========================================================================
    # 7. INPUT AND OUTPUT DESCRIPTION
    # =========================================================================
    add_heading_1(doc, "7. Input and Output Description")
    add_body_p(doc,
        "This section details the critical inputs required from users, the internal processing actions performed by the "
        "ParkFinity system, and the expected outputs and responses produced across all major operational workflows."
    )
    
    add_heading_2(doc, "7.1 Rider Module Inputs and Outputs")
    headers_io = ["Function / Operation", "Input Required from User", "System Processing / Action", "Expected Output / Result"]
    widths_io = [1.5, 1.6, 1.8, 1.6]
    
    rider_io_data = [
        [
            "User Registration",
            "Full Name, valid phone number (+880), password (min 8 chars), or Google OAuth token.",
            "Validates phone number format, checks for duplicate email/phone in PostgreSQL, hashes password with bcrypt.",
            "New rider account created in database; JWT authentication token issued; user redirected to role dashboard."
        ],
        [
            "Rider KYC Verification",
            "NID front & back photos, driving license photo.",
            "Uploads encrypted images to Supabase Storage, creates KYC record, sets verification status to 'pending'.",
            "Document submission confirmation message; 'Under Review' status badge displayed on user profile."
        ],
        [
            "Geospatial Map Search",
            "Current GPS coordinates, search query text (landmark/address), vehicle type filter.",
            "Executes PostGIS geospatial radius query, checks real-time slot availability, filters active spots.",
            "Interactive map updated with pins showing available parking spots with tariff tags and slot counters."
        ],
        [
            "AI Smart Parking Query",
            "Natural-language query string (e.g., 'Cheapest covered spot near NSTU with CCTV').",
            "Parses intent using text embedding & heuristic scoring across distance, price, vehicle type, and amenities.",
            "Ranked list of top recommendations with personalized AI explanation cards."
        ],
        [
            "Garage Vehicle Registration",
            "Vehicle type (Car/Bike/Pickup), manufacturer brand, model, license plate number.",
            "Validates license plate format against Bangladeshi BRTA standard, stores record in user's garage.",
            "Vehicle card added to Garage screen; available for one-touch selection during booking."
        ],
        [
            "Digital Wallet Top-up",
            "Top-up amount (BDT), preferred payment channel (bKash, Nagad, Card).",
            "Initializes SSLCommerz transaction session, verifies IPN callback upon payment completion, updates wallet balance.",
            "Digital wallet balance credited; transaction receipt logged in user's ledger."
        ],
        [
            "Spot Booking & Payment",
            "Selected spot ID, vehicle ID, start time, duration (hours/days), wallet debit confirmation.",
            "Locks parking slot atomically, calculates total fee with peak multiplier, debits rider wallet into platform escrow.",
            "Booking confirmation screen with active timer, QR code reference, and push notification sent to rider and host."
        ],
        [
            "QR Code Check-in",
            "Camera scan of host physical QR standee at parking gate.",
            "Decodes cryptographic spot payload, validates matching active booking, updates booking status to 'active'.",
            "Live parking session countdown timer initiates; notification sent to host confirming arrival."
        ],
        [
            "QR Code Checkout",
            "Camera scan of host QR code at conclusion of parking.",
            "Calculates total elapsed time, checks for overstay penalties, releases locked slot, settles host wallet balance.",
            "Session concluded receipt displayed; slot returned to available pool; review modal presented."
        ],
        [
            "Submit Spot Review",
            "Star rating (1 to 5 stars), written feedback text.",
            "Stores review record, recalculates listing's aggregate rating score.",
            "Review published on listing card; thank-you toast message displayed."
        ]
    ]
    tbl_r_io = doc.add_table(rows=1, cols=4)
    format_styled_table(tbl_r_io, widths_io, headers_io, rider_io_data)
    
    add_heading_2(doc, "7.2 Parking Space Owner Module Inputs and Outputs")
    owner_io_data = [
        [
            "Owner KYC Verification",
            "Government NID photos (front/back), property ownership deed / utility bill.",
            "Encrypts files in Supabase Storage, sets host KYC queue status to 'pending' for administrative review.",
            "KYC submitted confirmation; host permissions unlocked upon admin approval."
        ],
        [
            "Create Parking Listing",
            "Title, address, photos (min 3), video (15s), slot capacities (Car/Bike/Pickup), tariffs, amenities, schedule.",
            "Validates media uploads, verifies coordinates, creates listing and listing_slots records.",
            "Parking listing published; unique QR code standee generated; spot made discoverable on rider map."
        ],
        [
            "Download Spot QR Standee",
            "Host requests QR standee generation.",
            "Generates high-resolution printable QR code containing encrypted spot ID and 6-character short code.",
            "Downloadable PDF/image standee saved to device gallery for physical printing and mounting."
        ],
        [
            "Manage Slot Availability",
            "Toggle spot active/paused status, update operational hours or tariffs.",
            "Updates database listing record; immediately adjusts real-time visibility on rider map search.",
            "Updated status reflected instantly across the platform."
        ],
        [
            "Bank Revenue Withdrawal",
            "Withdrawal amount, bank name, branch, account number, routing number.",
            "Verifies available wallet balance, reserves withdrawal amount, logs pending payout request in admin queue.",
            "Withdrawal request submitted confirmation; balance debited from available funds; funds transferred upon admin approval."
        ]
    ]
    tbl_o_io = doc.add_table(rows=1, cols=4)
    format_styled_table(tbl_o_io, widths_io, headers_io, owner_io_data)

    add_heading_2(doc, "7.3 System Administrator Module Inputs and Outputs")
    admin_io_data = [
        [
            "Admin Authentication",
            "Administrator email address and secure password.",
            "Verifies bcrypt password hash against system admin role table, issues high-privilege administrative session token.",
            "Administrative Executive Dashboard unlocked with full platform governance tools."
        ],
        [
            "User Account Moderation",
            "User search query, selection of account, toggle suspend/reactivate.",
            "Updates user account status flag in database; invalidates active auth sessions if suspended.",
            "User status badge updated to 'Suspended' or 'Active'; access blocked or restored immediately."
        ],
        [
            "KYC Compliance Review",
            "Inspection of NID, license, or property deed; click 'Approve' or 'Reject' with reason.",
            "Updates user KYC status to 'verified' or 'rejected'; triggers push notification to user.",
            "User profile marked with verified badge; permissions granted or corrective prompt issued."
        ],
        [
            "Listing Moderation",
            "Listing search, inspect media and details; click 'Suspend'.",
            "Sets listing status to 'suspended'; hides listing from public map and search results.",
            "Listing removed from active map; notification sent to host explaining moderation rationale."
        ],
        [
            "Platform Policy Configuration",
            "Commission rate (0.0–1.0), peak multiplier (1.0–3.0), cancellation grace window.",
            "Updates platform_settings table in database; immediately alters real-time pricing calculations.",
            "System-wide pricing engine updated with new commission cut and dynamic multipliers."
        ],
        [
            "Revenue & Analytics Report",
            "Select date range (30d / 90d / Custom).",
            "Aggregates total gross volume, commission revenue, host payouts, and booking counts over selected timeframe.",
            "Interactive graphical revenue charts and summary KPI widgets rendered on dashboard."
        ]
    ]
    tbl_a_io = doc.add_table(rows=1, cols=4)
    format_styled_table(tbl_a_io, widths_io, headers_io, admin_io_data)

    # =========================================================================
    # 8. ERROR MESSAGES AND TROUBLESHOOTING
    # =========================================================================
    add_heading_1(doc, "8. Error Messages and Troubleshooting")
    add_body_p(doc,
        "This troubleshooting matrix provides comprehensive diagnostic procedures and corrective actions for common "
        "technical, environmental, and operational issues encountered by users."
    )
    
    headers_err = ["Problem / Error Scenario", "Possible Cause", "Corrective Solution / Action Steps"]
    widths_err = [1.8, 2.0, 2.7]
    error_data = [
        [
            "Unable to Log In / 'Invalid Credentials'",
            "1. Incorrect password or email entered.\n2. Account has not been registered yet.\n3. Account suspended by admin.",
            "1. Check credentials carefully for typos.\n2. Use 'Forgot Password' to reset password via email.\n3. If suspended, contact ParkFinity support."
        ],
        [
            "GPS Location Inaccurate / Map Not Centering",
            "1. Device GPS / Location sensor is disabled.\n2. App location permission not granted.\n3. Weak GPS satellite signal indoors.",
            "1. Enable GPS from device quick settings.\n2. Navigate to Phone Settings > Apps > ParkFinity > Permissions > Enable 'Precise Location'.\n3. Move closer to open sky or Wi-Fi."
        ],
        [
            "QR Code Scanner Fails to Scan Standee",
            "1. Camera lens smudged or dirty.\n2. Insufficient lighting or glare on printed standee.\n3. Camera permission denied.",
            "1. Clean camera lens.\n2. Turn on phone flashlight or type the 6-character spot short code manually.\n3. Grant camera permission in device settings."
        ],
        [
            "Payment Failed during Wallet Top-up",
            "1. Insufficient funds in MFS/bank account.\n2. SSLCommerz gateway session timeout.\n3. Network interruption during OTP submission.",
            "1. Check your bKash/Nagad/Bank balance.\n2. Retry transaction with a stable internet connection.\n3. If money was debited but not credited, contact support with SSLCommerz TranID."
        ],
        [
            "Booking Failed: 'Insufficient Wallet Balance'",
            "Available digital wallet balance is lower than the required parking reservation fee.",
            "Navigate to Profile > Wallet > tap '+ Add Funds' and top up the required amount via SSLCommerz, then retry booking."
        ],
        [
            "Booking Failed: 'Slot Already Occupied'",
            "Another driver completed booking for the last available slot simultaneously.",
            "Refresh the map or listing screen to view updated slot counts and select another nearby parking spot."
        ],
        [
            "KYC Upload Fails / 'File Size Too Large'",
            "1. Image file exceeds maximum allowed upload size (10 MB).\n2. Unsupported document format.\n3. Network timeout.",
            "1. Capture photos directly within the app camera or compress images before uploading.\n2. Ensure files are in JPG, PNG, or PDF format."
        ],
        [
            "KYC Verification Rejected by Admin",
            "1. Uploaded document blurry, cropped, or unreadable.\n2. Document name does not match user account name.\n3. Property ownership proof invalid.",
            "Check the specific rejection reason displayed on your profile, re-photograph clear, uncropped documents in good lighting, and resubmit."
        ],
        [
            "Push Notifications Not Received",
            "1. Notification permissions disabled.\n2. Battery optimization killing background services.",
            "1. Enable notifications in App Info.\n2. Set battery usage to 'Unrestricted' for ParkFinity in device battery settings."
        ],
        [
            "Listing Not Appearing on Public Map",
            "1. Host KYC verification is still pending.\n2. Listing status set to 'Paused' or 'Suspended'.\n3. All vehicle slots currently full.",
            "1. Complete KYC verification.\n2. Verify listing is set to 'Active' in My Listings.\n3. Free slots will automatically make it visible."
        ],
        [
            "Bank Withdrawal Request Pending",
            "Withdrawals undergo manual compliance auditing within 24–48 banking business hours.",
            "Ensure bank details (account number, branch, routing code) are accurate. If pending over 48 hours, contact finance support."
        ]
    ]
    tbl_err = doc.add_table(rows=1, cols=3)
    format_styled_table(tbl_err, widths_err, headers_err, error_data)

    # =========================================================================
    # 9. FREQUENTLY ASKED QUESTIONS (FAQ)
    # =========================================================================
    add_heading_1(doc, "9. Frequently Asked Questions (FAQ)")
    
    add_heading_2(doc, "9.1 General Platform Questions")
    add_bullet(doc, "Q1: What is ParkFinity and how does it work?",
        "Answer: ParkFinity is a smart parking optimization platform connecting drivers looking for parking with property owners who have available parking spaces. Drivers find and book spots via mobile app, and space owners earn money hosting parked vehicles.")
    add_bullet(doc, "Q2: Is my personal information and document data secure?",
        "Answer: Yes. All personal data, payment records, and uploaded KYC documents are encrypted using industry-standard SSL/TLS encryption in transit and AES-256 encryption at rest within Supabase Cloud Storage. Documents are accessible exclusively to authorized compliance officers for fraud verification.")
    add_bullet(doc, "Q3: Can one user operate as both a Rider and a Space Owner?",
        "Answer: Yes. A registered user can switch roles or submit KYC verification for both driver credentials and property hosting under a unified profile.")

    add_heading_2(doc, "9.2 Rider / Vehicle Driver Questions")
    add_bullet(doc, "Q4: How do I find the cheapest parking spot near my destination?",
        "Answer: You can use the search filter on the Listings screen to sort by price, or use 'Ask ParkFinity AI' to enter 'Find the cheapest parking near [destination]' to receive ranked recommendations.")
    add_bullet(doc, "Q5: What happens if I arrive late at the parking spot?",
        "Answer: Your reserved parking slot remains locked exclusively for your vehicle throughout your booked duration. However, the booked time window begins counting down from your scheduled start time.")
    add_bullet(doc, "Q6: What happens if I stay longer than my booked parking time (Overstaying)?",
        "Answer: If you need additional time, tap 'Extend' on the active session screen before your time expires. If you overstay without extending, an automatic overstay penalty rate will be charged to your wallet upon checkout.")
    add_bullet(doc, "Q7: Can I cancel a booking and receive a refund?",
        "Answer: Yes. If you cancel at least 30 minutes prior to your scheduled start time, a 100% refund is credited to your digital wallet immediately. Cancellations within 30 minutes may incur a nominal cancellation fee as per platform policy.")

    add_heading_2(doc, "9.3 Parking Space Owner Questions")
    add_bullet(doc, "Q8: What requirements must my parking space fulfill to be listed?",
        "Answer: The parking space must be legal, accessible via motorable road, unobstructed, and you must possess legal rights (ownership, tenancy, or authority) to host vehicles. Uploading 3 clear photos and a 15-second walkthrough video is mandatory.")
    add_bullet(doc, "Q9: How do drivers enter my parking space?",
        "Answer: You print and mount the dedicated Spot QR Code Standee at your parking gate. Arriving drivers scan the QR code using the ParkFinity app to validate their booking and check in.")
    add_bullet(doc, "Q10: How and when do I receive my rental earnings?",
        "Answer: Earnings are credited to your in-app wallet immediately after each completed booking. You can request a withdrawal to your registered Bangladeshi bank account at any time via the Wallet tab.")
    add_bullet(doc, "Q11: What platform fee does ParkFinity charge?",
        "Answer: ParkFinity retains a modest platform commission (configured between 5% and 10%) on completed transactions to maintain platform infrastructure, customer support, and insurance coverage.")

    # =========================================================================
    # 10. LOGOUT AND EXIT PROCEDURE
    # =========================================================================
    add_heading_1(doc, "10. Logout and Exit Procedure")
    add_body_p(doc,
        "To safeguard account security—particularly when utilizing shared or public devices—users should follow the "
        "recommended logout and application termination procedures."
    )
    add_heading_2(doc, "10.1 Mobile Application Logout")
    add_step(doc, 1, "Open Profile", "Tap the 'Profile' icon located on the far right of the bottom navigation bar.")
    add_step(doc, 2, "Select Logout", "Scroll to the bottom of the profile options and tap 'Log Out'.")
    add_step(doc, 3, "Confirm Action", "A confirmation dialogue will appear. Tap 'Yes, Log Out' to terminate your active session.")
    add_step(doc, 4, "Token Invalidation", "The application revokes the local JWT authentication token, clears secure cache buffers, and redirects to the Login screen.")

    add_heading_2(doc, "10.2 Admin Web Portal Logout")
    add_step(doc, 1, "Access Header Menu", "Click on your administrative profile avatar located at the top-right corner of the web navigation header.")
    add_step(doc, 2, "Click Sign Out", "Select 'Sign Out' from the dropdown menu.")
    add_step(doc, 3, "Session Termination", "The administrative JWT session is invalidated on the server and the browser is returned to the Admin Login screen.")

    # =========================================================================
    # 11. LIMITATIONS AND PRECAUTIONS
    # =========================================================================
    add_heading_1(doc, "11. Limitations and Precautions")
    add_body_p(doc,
        "Users and space hosts should remain cognizant of the following operational limitations and safety precautions:"
    )
    add_bullet(doc, "GPS Satellite Drift: ", 
        "In dense high-rise urban corridors, GPS accuracy may fluctuate by 5 to 15 meters. Always cross-reference the written spot address and landmark descriptions before parking.")
    add_bullet(doc, "Internet Connection Dependency: ", 
        "ParkFinity requires an active cellular data or Wi-Fi connection to process QR check-in, real-time timer updates, and payment validations. Ensure connectivity before entering basement garages.")
    add_bullet(doc, "Document Verification Latency: ", 
        "KYC document reviews are typically completed within 1 to 4 business hours. Users are advised to complete KYC verification ahead of time before urgent parking needs.")
    add_bullet(doc, "Host Facility Guidelines: ", 
        "Drivers must respect private property rules, observe parking slot boundaries, avoid blocking neighboring vehicles, and maintain clean parking habits.")
    add_bullet(doc, "Vehicle Height and Clearance: ", 
        "Ensure your vehicle dimensions (especially for SUVs, pickups, and roof-rack carriers) adhere to the spot's height and width specifications before booking.")

    # =========================================================================
    # 12. CONTACT AND SUPPORT INFORMATION
    # =========================================================================
    add_heading_1(doc, "12. Contact and Support Information")
    add_body_p(doc,
        "For customer support, bug reporting, billing inquiries, or technical assistance, contact the ParkFinity team "
        "through the official channels listed below:"
    )
    
    headers_sup = ["Channel / Department", "Contact Details", "Operating Hours"]
    widths_sup = [2.2, 2.8, 1.5]
    support_data = [
        ["General Customer Support Email", "support@parkfinity.com", "24/7 Response via Ticket"],
        ["Emergency Helpline (Bangladesh)", "+880 1845-503086 / +880 1886-044194", "8:00 AM – 10:00 PM (Daily)"],
        ["Billing & Host Payout Inquiries", "billing@parkfinity.com", "9:00 AM – 5:00 PM (Sun–Thu)"],
        ["Compliance & KYC Verification", "kyc@parkfinity.com", "9:00 AM – 6:00 PM (Daily)"],
        ["Academic & Institutional Inquiry", "Institute of Information Technology (IIT)\nNoakhali Science and Technology University (NSTU)\nNoakhali-3814, Bangladesh", "9:00 AM – 5:00 PM (Workdays)"]
    ]
    tbl_sup = doc.add_table(rows=1, cols=3)
    format_styled_table(tbl_sup, widths_sup, headers_sup, support_data)

    # =========================================================================
    # APPENDIX
    # =========================================================================
    add_heading_1(doc, "Appendix")
    
    add_heading_2(doc, "Appendix A: System Architecture & Technical Specifications")
    add_body_p(doc,
        "ParkFinity is built upon a modular, reactive, cloud-native architecture. The client tier comprises a Flutter "
        "mobile application targeting Android and iOS, alongside a Flutter Web administrative dashboard. The backend "
        "tier utilizes Supabase Cloud with PostgreSQL 15, PostGIS geospatial indexing, and PostgreSQL Row-Level Security "
        "(RLS) for multi-tenant data isolation. Third-party integrations include Google Maps Platform for navigation and "
        "geocoding, SSLCommerz for financial payment processing, Firebase Cloud Messaging (FCM) for push notifications, "
        "and Google ML Kit for on-device OCR and QR decoding."
    )
    
    add_heading_2(doc, "Appendix B: Database Entity Summary")
    add_bullet(doc, "users: ", "Stores core authentication profiles, contact info, assigned roles (rider, owner, admin), and KYC status.")
    add_bullet(doc, "kyc_documents: ", "Maintains uploaded NID, driving license, and property deed URLs with verification timestamps.")
    add_bullet(doc, "vehicles: ", "Stores rider registered vehicles, vehicle types, brands, models, and license plate numbers.")
    add_bullet(doc, "listings: ", "Stores parking spot titles, addresses, GPS coordinates (PostGIS points), amenities, and schedules.")
    add_bullet(doc, "listing_slots: ", "Maintains individual slot counts and pricing tariffs per vehicle type (Car, Motorcycle, Pickup).")
    add_bullet(doc, "bookings: ", "Manages reservation sessions, start/end timestamps, check-in status, total fees, and QR codes.")
    add_bullet(doc, "wallets & transactions: ", "Maintains user balances, SSLCommerz top-ups, booking deductions, and host payouts.")
    add_bullet(doc, "reviews: ", "Stores star ratings and user feedback for completed parking sessions.")
    add_bullet(doc, "platform_settings: ", "Maintains global commission percentages, dynamic multipliers, and refund rules.")

    # =========================================================================
    # SUPERVISOR APPROVAL PAGE
    # =========================================================================
    doc.add_page_break()
    
    p_ap_h = doc.add_paragraph()
    p_ap_h.paragraph_format.space_before = Pt(12)
    p_ap_h.paragraph_format.space_after = Pt(8)
    p_ap_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ap_h = p_ap_h.add_run("SUPERVISOR APPROVAL")
    r_ap_h.bold = True
    r_ap_h.font.name = "Times New Roman"
    r_ap_h.font.size = Pt(16)
    r_ap_h.font.color.rgb = COLOR_PRIMARY
    
    # Bordered Approval Container Box
    tbl_app = doc.add_table(rows=1, cols=1)
    tbl_app.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_app = tbl_app.cell(0, 0)
    c_app.width = Inches(6.5)
    set_cell_margins(c_app, top=160, bottom=160, left=180, right=180)
    set_cell_background(c_app, "FAFAFA")
    
    tcPr = c_app._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="12" w:space="0" w:color="1E3A8A"/>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="1E3A8A"/>'
        f'  <w:left w:val="single" w:sz="12" w:space="0" w:color="1E3A8A"/>'
        f'  <w:right w:val="single" w:sz="12" w:space="0" w:color="1E3A8A"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p_app_proj = c_app.paragraphs[0]
    p_app_proj.paragraph_format.space_before = Pt(0)
    p_app_proj.paragraph_format.space_after = Pt(8)
    p_app_proj.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_p1 = p_app_proj.add_run("Project Title: ")
    r_p1.bold = True
    r_p1.font.name = "Times New Roman"
    r_p1.font.size = Pt(11)
    r_p1.font.color.rgb = COLOR_PRIMARY
    r_p2 = p_app_proj.add_run("ParkFinity – A Smart Platform for Parking Space Optimization\n")
    r_p2.bold = True
    r_p2.font.name = "Times New Roman"
    r_p2.font.size = Pt(11)
    
    r_c1 = p_app_proj.add_run("Course: ")
    r_c1.bold = True
    r_c1.font.name = "Times New Roman"
    r_c1.font.size = Pt(10.5)
    r_c1.font.color.rgb = COLOR_PRIMARY
    r_c2 = p_app_proj.add_run("Software Project Lab-II (Course Code: SE 3112)\n")
    r_c2.font.name = "Times New Roman"
    r_c2.font.size = Pt(10.5)
    
    r_d1 = p_app_proj.add_run("Department/Institute: ")
    r_d1.bold = True
    r_d1.font.name = "Times New Roman"
    r_d1.font.size = Pt(10.5)
    r_d1.font.color.rgb = COLOR_PRIMARY
    r_d2 = p_app_proj.add_run("Institute of Information Technology (IIT), Noakhali Science and Technology University (NSTU)\n")
    r_d2.font.name = "Times New Roman"
    r_d2.font.size = Pt(10.5)
    
    # Students Table inside approval
    p_gm = c_app.add_paragraph()
    p_gm.paragraph_format.space_before = Pt(6)
    p_gm.paragraph_format.space_after = Pt(4)
    r_gm = p_gm.add_run("Group Members:")
    r_gm.bold = True
    r_gm.font.name = "Times New Roman"
    r_gm.font.size = Pt(11)
    r_gm.font.color.rgb = COLOR_PRIMARY
    
    students_app = [
        ("1", "Md. Shamsuddoha", "ASH2325005M"),
        ("2", "Eftekar Hossen", "ASH2325006M"),
        ("3", "Md. Majedur Rahman", "ASH2125021M")
    ]
    for num, sname, sid in students_app:
        p_srow = c_app.add_paragraph()
        p_srow.paragraph_format.space_before = Pt(1)
        p_srow.paragraph_format.space_after = Pt(2)
        p_srow.paragraph_format.left_indent = Inches(0.2)
        r_sn = p_srow.add_run(f"{num}. {sname} ")
        r_sn.bold = True
        r_sn.font.name = "Times New Roman"
        r_sn.font.size = Pt(10)
        r_si = p_srow.add_run(f"(ID: {sid})")
        r_si.font.name = "Times New Roman"
        r_si.font.size = Pt(10)
        r_si.font.color.rgb = COLOR_MUTED
        r_sig = p_srow.add_run("  ...................................................... Signature")
        r_sig.font.name = "Times New Roman"
        r_sig.font.size = Pt(9.5)
        r_sig.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        
    p_stmt = c_app.add_paragraph()
    p_stmt.paragraph_format.space_before = Pt(14)
    p_stmt.paragraph_format.space_after = Pt(12)
    p_stmt.paragraph_format.line_spacing = 1.15
    p_stmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_stmt = p_stmt.add_run(
        "I have reviewed the User Manual Report of the above-mentioned SPL-2 project titled 'ParkFinity – A Smart Platform "
        "for Parking Space Optimization'. The documentation conforms to the prescribed academic format, accurately details "
        "all system functionalities, user interfaces, and administrative workflows, and is hereby approved for final submission "
        "and SPL-2 presentation."
    )
    r_stmt.font.name = "Times New Roman"
    r_stmt.font.size = Pt(10)
    r_stmt.font.italic = True
    r_stmt.font.color.rgb = COLOR_DARK
    
    # Supervisor Signature Info
    p_sup_box = c_app.add_paragraph()
    p_sup_box.paragraph_format.space_before = Pt(14)
    p_sup_box.paragraph_format.space_after = Pt(2)
    p_sup_box.paragraph_format.line_spacing = 1.2
    
    r_sn1 = p_sup_box.add_run("Supervisor's Name: ")
    r_sn1.bold = True
    r_sn1.font.name = "Times New Roman"
    r_sn1.font.size = Pt(10.5)
    r_sn2 = p_sup_box.add_run("Md Hasan Imam\n")
    r_sn2.bold = True
    r_sn2.font.name = "Times New Roman"
    r_sn2.font.size = Pt(10.5)
    
    r_sd1 = p_sup_box.add_run("Designation: ")
    r_sd1.bold = True
    r_sd1.font.name = "Times New Roman"
    r_sd1.font.size = Pt(10)
    r_sd2 = p_sup_box.add_run("Assistant Professor\n")
    r_sd2.font.name = "Times New Roman"
    r_sd2.font.size = Pt(10)
    
    r_di1 = p_sup_box.add_run("Department/Institute: ")
    r_di1.bold = True
    r_di1.font.name = "Times New Roman"
    r_di1.font.size = Pt(10)
    r_di2 = p_sup_box.add_run("Institute of Information Technology (IIT)\n")
    r_di2.font.name = "Times New Roman"
    r_di2.font.size = Pt(10)
    
    r_un1 = p_sup_box.add_run("University: ")
    r_un1.bold = True
    r_un1.font.name = "Times New Roman"
    r_un1.font.size = Pt(10)
    r_un2 = p_sup_box.add_run("Noakhali Science and Technology University (NSTU)\n\n")
    r_un2.font.name = "Times New Roman"
    r_un2.font.size = Pt(10)
    
    r_sig_l = p_sup_box.add_run("Signature: _________________________________________           Date: ________________________\n")
    r_sig_l.bold = True
    r_sig_l.font.name = "Times New Roman"
    r_sig_l.font.size = Pt(10)
    r_sig_l.font.color.rgb = COLOR_PRIMARY
    


print("Sections 7 to 12 ready.")
