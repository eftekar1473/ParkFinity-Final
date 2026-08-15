import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def add_page_number_fields(run):
    """Inserts a dynamic Word PAGE field."""
    fldSimple = parse_xml(r'<w:fldSimple %s w:instr="PAGE"/>' % nsdecls('w'))
    run._r.append(fldSimple)

def add_numpages_fields(run):
    """Inserts a dynamic Word NUMPAGES field."""
    fldSimple = parse_xml(r'<w:fldSimple %s w:instr="NUMPAGES"/>' % nsdecls('w'))
    run._r.append(fldSimple)

def set_cell_margins(cell, top=120, bottom=120, left=180, right=180):
    """Set cell padding in dxa (twips)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, fill_hex):
    """Set background color of a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    """Set clean borders for table."""
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_callout_box(doc, title, text, icon="ℹ️", box_color="1E3A8A", bg_color="F0F4F8"):
    """Adds a stylish callout alert box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    set_cell_background(cell, bg_color)
    
    # Left border highlight
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{box_color}"/>'
        f'  <w:top w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    r_icon = p.add_run(f"{icon}  ")
    r_icon.font.name = "Segoe UI Emoji"
    r_icon.font.size = Pt(11)
    
    r_title = p.add_run(f"{title}: ")
    r_title.bold = True
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(10.5)
    r_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    r_text = p.add_run(text)
    r_text.font.name = "Times New Roman"
    r_text.font.size = Pt(10.5)
    r_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Add small spacing after table
    sp_p = doc.add_paragraph()
    sp_p.paragraph_format.space_before = Pt(0)
    sp_p.paragraph_format.space_after = Pt(6)

print("Base setup script template ready.")
