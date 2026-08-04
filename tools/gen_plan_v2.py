# -*- coding: utf-8 -*-
"""Generates Parkfinity_Implementation_Plan_v2.docx (Phases 11-19)."""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\SPL_2\Parkfinity_Implementation_Plan_v2.docx"

PURPLE = RGBColor(0x51, 0x2D, 0xA8)
GREY = RGBColor(0x55, 0x55, 0x55)
RED = RGBColor(0xB0, 0x00, 0x20)
GREEN = RGBColor(0x1B, 0x5E, 0x20)

doc = Document()

# ---------- base styles ----------
st = doc.styles['Normal']
st.font.name = 'Calibri'
st.font.size = Pt(10.5)
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.line_spacing = 1.12

for name, size, color in (('Heading 1', 20, PURPLE), ('Heading 2', 15, PURPLE), ('Heading 3', 12, RGBColor(0x33, 0x33, 0x33))):
    s = doc.styles[name]
    s.font.name = 'Calibri'
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = True


def h1(t):
    doc.add_page_break()
    return doc.add_heading(t, level=1)


def h1_nobreak(t):
    return doc.add_heading(t, level=1)


def h2(t):
    return doc.add_heading(t, level=2)


def h3(t):
    return doc.add_heading(t, level=3)


def p(t='', bold=False, italic=False, color=None, size=None, space=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space)
    r = par.add_run(t)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return par


def bullet(t, level=0):
    par = doc.add_paragraph(t, style='List Bullet')
    par.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    par.paragraph_format.space_after = Pt(2)
    return par


def num(t, level=0):
    par = doc.add_paragraph(t, style='List Number')
    par.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    par.paragraph_format.space_after = Pt(2)
    return par


def code(text):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Inches(0.22)
    par.paragraph_format.space_before = Pt(4)
    par.paragraph_format.space_after = Pt(8)
    pf = par._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), 'F4F2FA')
    pf.append(shd)
    r = par.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(8.5)
    rpr = r._element.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:ascii'), 'Consolas')
    rf.set(qn('w:hAnsi'), 'Consolas')
    rpr.append(rf)
    return par


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htxt in enumerate(headers):
        hdr[i].text = ''
        run = hdr[i].paragraphs[0].add_run(htxt)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(8.5)
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ==========================================================================
# TITLE
# ==========================================================================
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tp.add_run('PARKFINITY')
r.bold = True
r.font.size = Pt(40)
r.font.color.rgb = PURPLE

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('Smart Parking Marketplace')
r.font.size = Pt(15)
r.font.color.rgb = GREY

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('Implementation Plan v2  —  Phases 11 to 19')
r.bold = True
r.font.size = Pt(17)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sp.add_run('Defect remediation, QR access control, interval availability,\nrole-correct wallets, and production hardening')
r.font.size = Pt(11)
r.font.color.rgb = GREY

doc.add_paragraph()
table(
    ['Field', 'Value'],
    [
        ['Document', 'Parkfinity_Implementation_Plan_v2.docx'],
        ['Supersedes', 'Parkfinity_Implementation_Plan.docx (Phases 0-10, all complete)'],
        ['Date', '4 August 2026'],
        ['Client app', 'D:\\SPL_2\\parkfinity3  (Flutter 3.x, Riverpod 3, go_router 17)'],
        ['Admin app', 'D:\\SPL_2\\parkfinity_admin  (Flutter Web)'],
        ['Backend', 'Supabase — project ParkFinityDB, ref rkqduzjkkyplceipydir, ap-south-1, PG17'],
        ['Auth', 'Supabase GoTrue (email + Google ID-token)'],
        ['Push', 'Firebase Cloud Messaging HTTP v1'],
        ['Payments', 'SSLCommerz (sandbox) + internal wallet ledger'],
        ['Maps', 'Google Maps SDK + Places API (proxied)'],
        ['AI', 'Local weighted scoring engine + Groq llama-3.1-8b for narration'],
        ['Phases in scope', '11, 12, 13, 14, 15, 16, 17, 18, 19'],
    ],
    widths=[1.5, 4.7],
)

# ==========================================================================
h1('1. Executive Summary')

p('Phases 0 through 10 delivered a functional two-sided marketplace: schema, KYC, listings, '
  'the booking engine, overstay charging, refunds and payouts, reviews, FCM push, the admin '
  'panel, and localization scaffolding. Live testing then exposed a set of defects and design '
  'gaps that block a production launch. This document specifies the remediation and the new '
  'subsystems required to close them.')

h2('1.1 What is actually broken')
p('Every item below was traced to a specific line of code, not guessed. Root causes appear in '
  'Section 3.')
bullet('Sign-up is unreachable — the router evicts unauthenticated users from /register.')
bullet('Google sign-in presents the account chooser then fails silently.')
bullet('The profile screen renders a hardcoded fake identity ("Ahmed", ahmed@example.com, a pravatar image).')
bullet('Edit Profile is an empty placeholder widget and is not routed.')
bullet('NID and licence upload tiles are duplicated in the profile after the KYC gate already collected them.')
bullet('Booking cards are non-interactive; no detail view exists for either role.')
bullet('Dark mode has no effect because ~130 hardcoded colour literals override the theme.')
bullet('Bengali is incomplete — 23 missing ARB keys plus roughly 70 string literals never routed through l10n.')
bullet('Owners are shown an "Add Funds" button; owners must never fund the platform.')
bullet('Help Center and Privacy Policy tiles have empty tap handlers.')
bullet('Listing photo galleries are not viewable full-screen.')
bullet('Phone numbers are never collected, so the two parties cannot contact each other.')

h2('1.2 What is missing by design')
bullet('No proof-of-parking. The platform cannot tell whether a rider actually occupied the space.')
bullet('Availability is a plain counter. A one-hour booking blocks a slot indefinitely instead of freeing it at the end of the hour.')
bullet('Riders cannot choose an explicit start and end time.')
bullet('Peak-hour detection reads UTC clock fields, so the intended 08:00-10:00 and 17:00-20:00 Dhaka windows land on 02:00-04:00 and 11:00-14:00 Dhaka time.')
bullet('Owner earnings are credited only by complete_booking(), which nothing ever calls automatically.')
bullet('Map search has no autocomplete; GPS requires a manual button press; the rider is not visually distinguished from listings.')
bullet('Several notification types in the requirements have no producer.')

h2('1.3 Delivery shape')
table(
    ['Phase', 'Title', 'Layer', 'Migration', 'Risk'],
    [
        ['11', 'Auth & profile identity', 'Both', 'yes', 'Low'],
        ['12', 'QR check-in / check-out', 'Both', 'yes', 'High'],
        ['13', 'Interval-based availability', 'Backend-heavy', 'yes', 'High'],
        ['14', 'Booking details, contact, navigation', 'Frontend', 'no', 'Low'],
        ['15', 'Wallet role split & payout automation', 'Both', 'yes', 'Medium'],
        ['16', 'Explore map upgrade', 'Both', 'no (edge fn)', 'Medium'],
        ['17', 'Theme, i18n, static pages', 'Frontend', 'no', 'Low'],
        ['18', 'Notification completeness', 'Backend-heavy', 'yes', 'Medium'],
        ['19', 'Hardening, admin catch-up, QA', 'Both', 'yes', 'Medium'],
    ],
    widths=[0.6, 2.4, 1.1, 0.9, 0.7],
)
p('Phases 12 and 13 are coupled: check-out cannot finalise a booking until availability is '
  'interval-aware, and interval availability cannot release capacity correctly until check-out '
  'writes actual_end_time. They are specified separately but shipped as one migration pair.',
  italic=True, color=GREY)

# ==========================================================================
h1('2. Current System — Verified Baseline')

h2('2.1 Client inventory')
p('66 Dart files under parkfinity3/lib, feature-first, one folder per bounded context.')
table(
    ['Layer', 'Contents'],
    [
        ['core/', 'router (go_router + auth redirect), theme, services (PushService), controllers (settings), utils'],
        ['features/auth/', 'auth_repository, kyc_repository, auth_controller, login / register / role / KYC / splash / onboarding screens'],
        ['features/owner/', 'listing_model, listings_repository, withdrawal_repository, controllers, 7 screens, listing_form_fields widgets'],
        ['features/rider/', 'ai_recommendation_service, listing_filter, vehicle_model, reviews & vehicles repos, 4 controllers, 8 screens, filter_sheet'],
        ['features/wallet/', 'wallet_provider, wallet_screen, ssl_webview_screen'],
        ['features/shared/', 'document_verification_service, notification_service, profiles_repository, profile / notifications / edit_profile screens, review_sheet'],
        ['shared/', 'booking_model, bookings_repository'],
        ['l10n/', 'app_en.arb (290 keys), app_bn.arb (267 keys), generated delegates'],
    ],
    widths=[1.3, 4.9],
)

h2('2.2 Backend inventory')
p('14 migrations applied, 9 edge functions deployed.')
table(
    ['Edge function', 'Purpose', 'Status after v2'],
    [
        ['create-booking', 'Validate, lock slot, price, charge wallet, insert, notify', 'Rewritten in Phase 13'],
        ['extend-booking', 'Price extra units, guard next booking, push end_time', 'Rewritten in Phase 13'],
        ['overstay-checker', 'Thin wrapper over process_overstays()', 'Reworked in Phase 12'],
        ['sslcommerz-init', 'Start hosted payment session', 'Unchanged'],
        ['sslcommerz-webhook', 'Validate val_id, mark payment complete', 'Unchanged'],
        ['send-push', 'FCM HTTP v1 dispatch with cached OAuth token', 'Unchanged'],
        ['ai-recommendations', 'Groq narration over local ranking', 'Unchanged'],
        ['sync-profile', 'Idempotent profile upsert', 'Superseded by DB trigger, Phase 11'],
        ['_shared/pricing.ts', 'Base x peak x weekend x demand, commission split', 'Timezone fix, Phase 13'],
        ['places-autocomplete', 'NEW — proxy Google Places with a server-held key', 'Added in Phase 16'],
    ],
    widths=[1.5, 3.2, 1.5],
)

h2('2.3 Money flow as it exists today')
p('Verified correct in principle, incomplete in execution.')
code(
"""rider wallet  --(create-booking: full total debited)-->  platform float
                                                             |
                             commission_amount  <-------------+
                             owner_earnings     <-------------+
                                                             |
owner wallet  <--(complete_booking: owner_earnings credited)--+
                                                             |
bank account  <--(request_withdrawal -> admin approve)--------+""")
p('The split logic, the ledger rows, and the withdrawal hold/refund are all implemented and '
  'tested. The single break is that complete_booking() has no automatic caller — bookings sit '
  'in Confirmed forever and owners are never paid. Phase 12 (QR check-out) and Phase 15 '
  '(sweeper fallback) close that.', color=RED)

# ==========================================================================
h1('3. Defect Register — Root Cause Analysis')

p('Each defect below is traced to code. "Fix" states the intended change; the phase column '
  'says where it lands.')

h2('3.1 Authentication')

h3('D-01  Sign-up redirects back to login')
p('Reported: "Sign up also do not work, if I click sign up, it redirect to the login page again."')
p('Root cause — app_router.dart lines 57 and 73:', bold=True)
code(
"""final isLoggingIn = state.matchedLocation == '/login'
                 || state.matchedLocation == '/register';
...
if (!isAuth) {
  return isSplash || isLoggingIn ? '/login' : (isSplash ? null : '/login');
}""")
p('/register is classified as a "logging in" location, and the unauthenticated branch returns '
  "'/login' for it unconditionally. Navigating to /register therefore triggers an immediate "
  'redirect to /login. The ternary is also dead code — every branch yields the same value.')
p('Fix (Phase 11): allow unauthenticated users to remain on any public auth route.', bold=True)
code(
"""const publicRoutes = {'/login', '/register', '/onboarding', '/forgot_password'};
if (!isAuth) {
  return publicRoutes.contains(loc) ? null : '/login';
}""")

h3('D-02  Google sign-in stalls after the account chooser')
p('Reported: "Continue with google is not working, it shows all the gmail but does not login if I click."')
p('The account picker rendering proves the Android OAuth client resolves. The failure is after '
  'token exchange. Three candidate causes, all checked in Phase 11:')
num('serverClientId must be the Web OAuth client ID, not the Android one. If GOOGLE_WEB_CLIENT_ID '
    'in .env holds the Android client, googleAuth.idToken returns null and auth_repository.dart '
    'line 65 throws "No ID token from Google."')
num('The debug keystore SHA-1 must be registered on the Android OAuth client in Google Cloud '
    'Console. Without it the picker still shows but the ID token is refused.')
num('Supabase Dashboard -> Authentication -> Providers -> Google must have the same Web client ID '
    'listed under Authorized Client IDs, otherwise signInWithIdToken rejects a structurally valid token.')
p('Additional defect: the error surfaces as a SnackBar that is dismissed by the redirect, so the '
  'user perceives silence. Phase 11 replaces this with a persistent inline error banner plus a '
  'debug-mode diagnostic dump.')

h3('D-03  No phone number captured')
p('profiles.phone_number exists in the initial schema and is never written. Both roles need it '
  'for the click-to-call requirement.')
p('Fix (Phase 11): mandatory phone field on the registration form and on the KYC gate for '
  'Google-origin accounts, validated against the Bangladeshi mobile format ^(?:\\+?88)?01[3-9]\\d{8}$, '
  'stored normalised as +8801XXXXXXXXX.')

h2('3.2 Profile')

h3('D-04  Hardcoded fake identity')
p('profile_screen.dart lines 146-167 literally hardcode the avatar, name, email, and rating:')
code(
"""const CircleAvatar(radius: 50,
  backgroundImage: NetworkImage('https://i.pravatar.cc/150?img=11')),
const Text('Ahmed', ...),
Text('ahmed@example.com', ...),
const Text('4.8 Rating', ...)""")
p('Fix (Phase 11): a currentProfileProvider streams the caller\'s profiles row; the header binds '
  'to full_name, email, avatar_url, and a real aggregate rating drawn from the reviews table. '
  'A shimmer placeholder covers the loading state and initials render when avatar_url is null.')

h3('D-05  Edit Profile is a stub')
p('edit_profile_screen.dart is a 15-line placeholder returning a Text widget, and no route points '
  'at it. profile_screen.dart line 203 wires onTap to an empty closure.')
p('Fix (Phase 11): full implementation with avatar picker, crop-to-square, upload to the Supabase '
  'avatars bucket, and an update_my_profile RPC.')

h3('D-06  Duplicate KYC uploads in profile')
p('profile_screen.dart lines 219-228 expose "Upload NID" and "Upload Driving License" tiles, but '
  'the KYC gate (Phase 1) already blocks app entry until both documents are submitted. The tiles '
  'also call verifyNid() on a licence image, which mis-validates.')
p('Fix (Phase 11): remove both tiles. Replace with a read-only "Verification status" tile showing '
  'kyc_status and thumbnails of what was submitted, plus a "Re-submit documents" action that is '
  'visible only when kyc_status is rejected.')

h2('3.3 Bookings')

h3('D-07  Booking cards are not tappable')
p('Reported for both roles. rider_booking_history_screen.dart builds a plain Container '
  '(_buildBookingCard, line 128) with no gesture handler; owner_booking_history_screen.dart is '
  'structurally identical. No detail route exists.')
p('Fix (Phase 14): a shared BookingDetailScreen with role-aware sections, reached from both lists.')

h3('D-08  Status classification is wrong')
p('rider_booking_history_screen.dart line 44 filters the "Upcoming & Active" tab on '
  "status == 'Pending' only:")
code(
"""final activeBookings = bookings.where(
    (b) => b.status == 'Pending' && b.endTime.isAfter(now)).toList();""")
p("Since Phase 4, instant bookings are inserted as Confirmed, not Pending. Confirmed, Active, and "
  "Overstayed bookings therefore fall into the Past tab. active_parking_screen.dart was already "
  "corrected in Phase 4 but the history screen was not.")
p('Fix (Phase 14): a single BookingStatusPolicy helper shared by every screen, so classification '
  'lives in one place.')

h2('3.4 Wallet')

h3('D-09  Owners see "Add Funds"')
p('app_router.dart line 246 maps /owner/wallet to the same WalletScreen as /rider/wallet, and '
  'that screen unconditionally renders the Add Funds button (wallet_screen.dart line 189).')
p('Fix (Phase 15): split into RiderWalletScreen and OwnerEarningsScreen sharing a common ledger '
  'widget. The owner variant shows lifetime earnings, pending clearance, available balance, and '
  'a Withdraw action — never a top-up.')

h2('3.5 Presentation')

h3('D-10  Dark mode does not apply')
p('Reported: "Dark theme not working, it only make the fonts dimmer, not the page." The themes '
  'themselves are correct — app_theme.dart builds proper light and dark ColorSchemes and main.dart '
  'wires themeMode. The problem is that screens bypass the theme. Measured occurrences of '
  'Colors.white / Colors.black / Colors.grey[n] / Colors.grey.shade[n]:')
table(
    ['File', 'Count'],
    [
        ['rider/.../listing_details_screen.dart', 16],
        ['rider/.../checkout_screen.dart', 13],
        ['shared/.../profile_screen.dart', 12],
        ['rider/.../my_vehicles_screen.dart', 10],
        ['rider/.../explore_map_screen.dart', 10],
        ['owner/.../add_listing_screen.dart', 9],
        ['owner/.../my_listings_screen.dart', 8],
        ['owner/.../owner_dashboard_screen.dart', 7],
        ['6 further screens', '6 each'],
        ['9 remaining files', '2-5 each'],
        ['TOTAL', '~130 across 23 files'],
    ],
    widths=[4.0, 1.0],
)
p('Because Scaffold backgrounds are pinned to Colors.white / Colors.grey[50] while text colour '
  'still resolves from the dark ColorScheme, the result is exactly what was reported: light pages '
  'with washed-out text. This is also an accessibility failure — contrast falls below WCAG AA.')
p('Fix (Phase 17): mechanical substitution to Theme.of(context).colorScheme tokens, plus a CI lint '
  'that fails the build on new hardcoded colour literals in lib/features.')

h3('D-11  Bengali coverage is incomplete')
p('app_en.arb holds 290 keys; app_bn.arb holds 267. Beyond the 23 missing translations, roughly 70 '
  'string literals were never extracted at all:')
table(
    ['File', 'Untranslated literals'],
    [
        ['owner/.../edit_listing_screen.dart', 18],
        ['shared/.../profile_screen.dart', 14],
        ['wallet/.../wallet_screen.dart', 13],
        ['owner/.../listing_form_fields.dart', 8],
        ['rider/.../filter_sheet.dart', 5],
        ['9 further files', '1-3 each'],
    ],
    widths=[4.0, 1.6],
)
p('Fix (Phase 17): extract every literal, reach full key parity, and add an arb_parity test that '
  'fails when the two files diverge.')

h3('D-12  Photo gallery not viewable')
p('Reported: "Rider cannot see the multiple pictures of the parking spot when click in the picture." '
  'listing_details_screen.dart renders photos in a PageView with no tap target.')
p('Fix (Phase 16): a full-screen PhotoGalleryScreen with pinch-zoom, swipe, a counter, and a '
  'Hero transition from the card.')

h3('D-13  Help Center and Privacy Policy are dead')
p('profile_screen.dart lines 285 and 290 both pass empty closures.')
p('Fix (Phase 17): two localized static content screens with real placeholder copy in English and Bengali.')

# ==========================================================================
h1('4. Architecture Decisions')

p('Six decisions govern the new subsystems. Each records the alternatives that were rejected, '
  'so the reasoning survives.')

h2('ADR-01  QR-based physical access control')
p('Context', bold=True)
p('The platform has no way to know whether a rider physically occupied a space. Overstay charging '
  'currently fires off wall-clock time alone, so a rider who never arrived is still billed a penalty, '
  'and an owner whose space was never used still gets paid.')
p('Decision', bold=True)
p('Every listing carries an immutable, unguessable QR token. The owner prints it and mounts it at '
  'the space. The rider scans once to start and once to stop. The scan is the trusted signal.')
table(
    ['Property', 'Choice', 'Reason'],
    [
        ['Scope', 'Per listing, not per owner', 'One owner may hold spaces at different addresses'],
        ['Token', 'gen_random_uuid() stored in listings.qr_token', 'Unguessable; rotatable without touching the primary key'],
        ['Payload', 'parkfinity://checkin?t=<uuid>', 'Custom scheme, no PII, works offline'],
        ['Fallback', '6-character alphanumeric code printed under the QR', 'Broken camera, dirty print, poor light'],
        ['Authority', 'Server-side RPC decides; the app only transports the token', 'The scan cannot be forged client-side'],
        ['Geofence', 'Optional, 150 m advisory radius', 'Warns rather than blocks — GPS drifts in car parks'],
    ],
    widths=[1.0, 2.4, 2.8],
)
p('Rejected alternatives', bold=True)
bullet('Bluetooth beacons — hardware cost per space, unsuitable for a marketplace of private owners.')
bullet('ANPR cameras — capital expense and privacy exposure far beyond project scope.')
bullet('Pure GPS geofencing — multi-storey and basement car parks routinely lose fix; a hard geofence would lock legitimate riders out.')
bullet('Owner-confirms-manually — reintroduces the manual process the product exists to remove, and creates a dispute surface.')

h2('ADR-02  Interval-based availability replaces counters')
p('Context', bold=True)
p('Phase 0 modelled availability as slot_capacity and slot_available JSONB counters, decremented '
  'by book_slot() and restored by release_slot(). A one-hour booking therefore consumes a slot from '
  'the moment of purchase until release, so nobody can book that space for any other hour of the '
  'day. This directly contradicts the requirement that a spot booked 14:00-15:00 be freely bookable '
  'for 15:00 onward, and for all other slots concurrently.')
p('Decision', bold=True)
p('Capacity stays a per-vehicle-type map. Occupancy becomes a query over overlapping time ranges.')
code(
"""available(listing, vtype, [start, end))
  = slot_capacity->>vtype
  - COALESCE((
      SELECT SUM(slot_qty) FROM bookings
       WHERE listing_id = listing
         AND vehicle_type = vtype
         AND status IN ('Pending','Confirmed','Active','Overstayed')
         AND tstzrange(start_time, COALESCE(actual_end_time, end_time), '[)')
             && tstzrange(start, end, '[)')
    ), 0)""")
p('Concurrency', bold=True)
p('Read-then-write is not safe under load. Every reservation takes a transaction-scoped advisory '
  'lock keyed on the listing before it computes availability, so two riders competing for the last '
  'slot serialise:')
code("PERFORM pg_advisory_xact_lock(hashtextextended(p_listing::text, 0));")
p('This is cheaper than SERIALIZABLE isolation and contends only per listing, not globally.')
p('Compatibility', bold=True)
p('listings.slot_available is retained as a denormalised "free right now" figure, recomputed by a '
  'trigger, so the map and existing filter code keep working unchanged. It becomes derived data, '
  'never a source of truth.')
p('Rejected alternatives', bold=True)
bullet('PostgreSQL EXCLUDE constraint with a GIST index — enforces at most one booking per range and cannot express "N of this type may overlap".')
bullet('Materialised slot-instance rows, one per hour per slot — a monthly booking on a 20-slot listing would generate 14,400 rows.')
bullet('Application-layer locking in the edge function — Deno instances are horizontally scaled and share no lock state.')

h2('ADR-03  Peak pricing is timezone-correct and configurable')
p('Context', bold=True)
p('_shared/pricing.ts reads start.getUTCHours() and start.getUTCDay(). Bangladesh is UTC+6 with no '
  'daylight saving, so the intended local windows are displaced by six hours:')
table(
    ['Intent (Asia/Dhaka)', 'What the code actually matches', 'Effect'],
    [
        ['Morning peak 08:00-10:00', '02:00-04:00 Dhaka', 'Surge applied to the emptiest hours'],
        ['Evening peak 17:00-20:00', '11:00-14:00 Dhaka', 'Surge on mid-day, none on the real rush'],
        ['Weekend Fri-Sat', 'Fri 06:00 - Sun 06:00 Dhaka', 'Thursday evening under-priced, Sunday morning over-priced'],
    ],
    widths=[1.8, 2.2, 2.2],
)
p('Decision', bold=True)
p('Introduce a single offset constant applied before any calendar field is read, and move the '
  'windows out of code into platform_settings so the admin can retune them without a redeploy.')
code(
"""const DHAKA_OFFSET_MIN = 360;   // UTC+6, no DST
function local(d: Date): Date {
  return new Date(d.getTime() + DHAKA_OFFSET_MIN * 60_000);
}
// peak_windows  jsonb: [{"from":8,"to":10},{"from":17,"to":20}]
// weekend_days  int[]: [5,6]   -- Fri, Sat""")
p('Peak resolution order (first match wins)', bold=True)
num('Learned demand — three or more bookings started in the same hour bucket over the trailing seven days (listing_hour_demand RPC). Multiplier scales with count, capped at peak_multiplier.')
num('Configured windows — platform_settings.peak_windows.')
num('Neutral — multiplier 1.0.')
p('Occupancy surge (demandMultiplier) stays multiplicative on top, and the composed factor is '
  'clamped so no booking can exceed 2.5x its base rate. The full breakdown is returned to the '
  'client and rendered line by line at checkout, so the rider sees exactly why the price moved.')

h2('ADR-04  Wallet semantics differ by role')
p('Context', bold=True)
p('Reported: "Why the owner need to have to add funds? Owner will just earn money." Correct — a '
  'supply-side participant funding the platform is a category error, and it invites a fraud '
  'pattern where a compromised owner account is used as a money-in channel.')
p('Decision', bold=True)
table(
    ['Capability', 'Rider', 'Owner', 'Admin'],
    [
        ['Top up via SSLCommerz', 'yes', 'NO', 'no'],
        ['Pay for a booking', 'yes', 'no', 'no'],
        ['Receive earnings', 'no', 'yes', 'no'],
        ['Receive refunds', 'yes', 'no', 'no'],
        ['Request withdrawal', 'no', 'yes', 'no'],
        ['Approve withdrawal', 'no', 'no', 'yes'],
    ],
    widths=[2.2, 1.2, 1.2, 1.2],
)
p('One profiles.wallet_balance column continues to serve both roles — the ledger in transactions '
  'already distinguishes intent by type. The restriction is enforced in two places: the UI never '
  'renders the affordance, and a database CHECK rejects a deposit transaction whose owner holds '
  'the Owner role.')
p('The owner surface additionally exposes three derived figures that the raw balance cannot '
  'express: lifetime gross, pending clearance (Confirmed and Active bookings not yet checked out), '
  'and available-to-withdraw (balance minus pending withdrawal holds).')

h2('ADR-05  Places Autocomplete is proxied, never called from the device')
p('Context', bold=True)
p('Search requires suggestions as the user types. The Google Maps key is currently embedded in '
  'AndroidManifest.xml with an Android application restriction. The Places Web Service does not '
  'accept application-restricted keys, so a second, unrestricted key would be required — and an '
  'unrestricted key shipped in an APK is extractable and billable by anyone who finds it.')
p('Decision', bold=True)
p('A places-autocomplete edge function holds the key as a Supabase secret. The client sends the '
  'query, a session token, and its coordinates; the function calls Google and returns only the '
  'fields the UI needs.')
code(
"""client -> POST /functions/v1/places-autocomplete
           { q, session_token, lat, lng }
edge   -> GET  maps.googleapis.com/maps/api/place/autocomplete/json
           ?input=..&components=country:bd&location=lat,lng&radius=50000
           &sessiontoken=..&key=<secret>
edge   -> [{ place_id, main_text, secondary_text }]""")
bullet('Session tokens group keystrokes into one billable session, cutting Places cost by roughly an order of magnitude.')
bullet('Results are biased to the rider\'s position and restricted to Bangladesh.')
bullet('The function debounces and caches identical queries for 60 seconds in memory.')
bullet('Key rotation is a secret update, not an app release.')

h2('ADR-06  Scheduled notifications run in-database')
p('Context', bold=True)
p('"Time ending" reminders and overstay warnings must fire at a wall-clock moment, not in response '
  'to a user action. Phase 8 established that every row inserted into notifications is pushed '
  'automatically by the dispatch_push trigger.')
p('Decision', bold=True)
p('A single pg_cron job every five minutes calls notification_sweep(), a plpgsql function that '
  'inserts the due notification rows. The existing trigger delivers them. No HTTP hop, no service '
  'key in transit, no external scheduler, no cost.')
p('Idempotency is guaranteed by a partial unique index on (booking_id, type) for the sweep-generated '
  'types, so a re-run or an overlapping tick cannot double-notify.')
p('Rejected: an external cron service calling an edge function (adds a dependency, a secret, and a '
  'failure mode for zero benefit), and client-side local notifications (do not fire when the app '
  'is terminated).')

# ==========================================================================
h1('5. Data Model Changes')

h2('5.1 New and altered columns')
table(
    ['Table', 'Column', 'Type', 'Purpose', 'Phase'],
    [
        ['profiles', 'phone_number', 'varchar(20) NOT NULL after backfill', 'Click-to-call; normalised +8801XXXXXXXXX', '11'],
        ['profiles', 'avatar_url', 'text (existing, now written)', 'Real profile picture', '11'],
        ['profiles', 'notification_prefs', 'jsonb DEFAULT ...', 'Per-channel and per-type opt-out', '18'],
        ['listings', 'qr_token', 'uuid UNIQUE DEFAULT gen_random_uuid()', 'Physical access token', '12'],
        ['listings', 'qr_short_code', 'varchar(6) UNIQUE', 'Manual fallback for the QR', '12'],
        ['listings', 'qr_rotated_at', 'timestamptz', 'Audit of token rotation', '12'],
        ['listings', 'contact_phone', 'varchar(20)', 'Overrides owner phone for this space', '11'],
        ['bookings', 'checked_in_at', 'timestamptz', 'First successful scan', '12'],
        ['bookings', 'checked_out_at', 'timestamptz', 'Second successful scan', '12'],
        ['bookings', 'check_in_method', "text CHECK (in 'qr','code','auto','admin')", 'Audit of how entry happened', '12'],
        ['bookings', 'no_show', 'boolean DEFAULT false', 'Never scanned in before the grace window closed', '12'],
        ['bookings', 'price_breakdown', 'jsonb', 'Frozen multipliers for dispute resolution', '13'],
        ['bookings', 'idempotency_key', 'text UNIQUE', 'Blocks duplicate submissions', '19'],
        ['platform_settings', 'peak_windows', 'jsonb', 'Admin-tunable peak hours', '13'],
        ['platform_settings', 'weekend_days', 'int[]', 'Admin-tunable weekend', '13'],
        ['platform_settings', 'max_price_multiplier', 'numeric DEFAULT 2.5', 'Composed surge ceiling', '13'],
        ['platform_settings', 'checkin_grace_minutes', 'int DEFAULT 15', 'Early-scan tolerance', '12'],
        ['platform_settings', 'no_show_minutes', 'int DEFAULT 30', 'No-show cutoff after start', '12'],
        ['platform_settings', 'reminder_minutes', 'int DEFAULT 30', 'Time-ending reminder lead', '18'],
    ],
    widths=[1.0, 1.4, 1.5, 1.8, 0.5],
)

h2('5.2 New tables')
table(
    ['Table', 'Purpose', 'Phase'],
    [
        ['scan_events', 'Append-only log of every scan attempt including rejections — dispute evidence and abuse detection', '12'],
        ['static_pages', 'Help Center and Privacy Policy content, keyed by slug and locale, editable from the admin panel', '17'],
        ['audit_log', 'Admin mutations: who changed what, when, from what to what', '19'],
    ],
    widths=[1.3, 4.4, 0.5],
)

h2('5.3 New and replaced database functions')
table(
    ['Function', 'Signature and behaviour', 'Phase'],
    [
        ['available_between', 'available_between(p_listing uuid, p_vtype text, p_start timestamptz, p_end timestamptz) -> int. Core interval query.', '13'],
        ['listing_availability', 'listing_availability(p_listing uuid, p_start, p_end) -> jsonb. Per-type map for the details screen.', '13'],
        ['reserve_interval', 'reserve_interval(...) -> (ok bool, msg text, available int). Advisory-locked; replaces book_slot.', '13'],
        ['recompute_slot_available', 'Trigger function keeping listings.slot_available current for map rendering.', '13'],
        ['check_in', 'check_in(p_token uuid, p_short_code text, p_booking uuid, p_lat, p_lng) -> jsonb. Validates and activates.', '12'],
        ['check_out', 'check_out(p_token uuid, p_booking uuid) -> jsonb. Settles overstay, completes, pays the owner.', '12'],
        ['process_overstays', 'Reworked: accrues penalty without releasing capacity until check-out.', '12'],
        ['mark_no_shows', 'Cancels bookings never scanned in, applies the no-show policy, frees the interval.', '12'],
        ['notification_sweep', 'Emits reminder, overstay-warning, review-prompt, and payout rows. pg_cron every 5 min.', '18'],
        ['settle_stale_bookings', 'Auto-completes bookings past end_time that were checked in but never checked out.', '15'],
        ['owner_earnings_summary', 'owner_earnings_summary(p_owner uuid) -> jsonb. Lifetime, pending, available.', '15'],
        ['update_my_profile', 'Guarded self-update of name, phone, and avatar. Email and role are immutable here.', '11'],
        ['my_profile', 'Returns the caller profile joined with an aggregate rating.', '11'],
    ],
    widths=[1.3, 4.4, 0.5],
)

h2('5.4 Indexes')
code(
"""-- interval overlap probe (the hot path of every availability check)
CREATE INDEX idx_bookings_listing_window
  ON bookings (listing_id, vehicle_type, start_time, end_time)
  WHERE status IN ('Pending','Confirmed','Active','Overstayed');

-- QR resolution
CREATE UNIQUE INDEX idx_listings_qr_token      ON listings (qr_token);
CREATE UNIQUE INDEX idx_listings_qr_short_code ON listings (qr_short_code);

-- notification idempotency (sweep-generated types only)
CREATE UNIQUE INDEX idx_notif_once ON notifications (data->>'booking_id', type)
  WHERE type IN ('booking_reminder','overstay_warning','review_prompt');

-- scan forensics
CREATE INDEX idx_scan_events_booking ON scan_events (booking_id, created_at DESC);""")

# ==========================================================================
h1('6. Phase 11 — Authentication and Profile Identity')

p('Closes D-01, D-02, D-03, D-04, D-05, D-06. No behavioural risk to money or bookings.')

h2('6.1 Backend')
p('Migration 20260804100000_phase11_profile.sql', bold=True)
num('Add profiles.phone_number constraint and a normalisation trigger; backfill existing rows to NULL and allow a one-time in-app completion prompt.')
num('Add listings.contact_phone.')
num('Recreate the handle_new_user trigger as an idempotent migration (currently only in _archive_applied, so a fresh environment would have no profile rows). It must read full_name, phone, avatar_url, and role from raw_user_meta_data and cope with Google sign-ups that supply name and picture but no phone.')
num('Create the avatars storage bucket, public read, owner-scoped write.')
num('Add update_my_profile(p_full_name, p_phone, p_avatar_url) SECURITY DEFINER — validates the phone format, rejects attempts to change email or role, and mirrors full_name and avatar_url into auth.users metadata so the JWT stays consistent.')
num('Add my_profile() returning the profile plus avg_rating and review_count from the reviews table.')
num('Deprecate the sync-profile edge function; the trigger supersedes it.')

h2('6.2 Frontend')
table(
    ['File', 'Action'],
    [
        ['core/router/app_router.dart', 'Rewrite the unauthenticated branch (D-01); add /rider/profile/edit, /owner/profile/edit, /help, /privacy'],
        ['features/auth/.../register_screen.dart', 'Add phone field with live validation; add a confirm-password field; convert to a Form with validators; replace transient SnackBar errors with an inline banner'],
        ['features/auth/.../login_screen.dart', 'Same inline error treatment; add "Forgot password"'],
        ['features/auth/data/auth_repository.dart', 'Pass phone into signUp metadata; wrap Google sign-in in a typed GoogleAuthFailure carrying the underlying cause; add resetPassword'],
        ['features/shared/data/profiles_repository.dart', 'Add MyProfile model, currentProfileProvider (stream), updateProfile, uploadAvatar'],
        ['features/shared/.../profile_screen.dart', 'Bind the header to real data (D-04); delete the two KYC tiles (D-06); add a verification-status tile; route Edit Profile, Help, Privacy'],
        ['features/shared/.../edit_profile_screen.dart', 'Full implementation (D-05): name, phone, avatar picker with square crop, upload progress, optimistic update, error rollback'],
        ['features/shared/.../widgets/avatar_widget.dart', 'NEW — network avatar with initials fallback and shimmer loading, reused in eight places'],
    ],
    widths=[2.3, 3.9],
)

h2('6.3 Google sign-in remediation runbook')
num('Print the release and debug SHA-1 with gradlew signingReport.')
num('In Google Cloud Console, confirm an Android OAuth client exists for com.parkfinity.parkfinity with both fingerprints, and a separate Web OAuth client.')
num('Set GOOGLE_WEB_CLIENT_ID in .env to the Web client ID (ending .apps.googleusercontent.com).')
num('In Supabase Dashboard -> Authentication -> Providers -> Google, enable the provider and paste the same Web client ID into Authorized Client IDs.')
num('Verify end to end: signInWithGoogle must return a non-null idToken and a session, and the profile row must materialise from the trigger.')
p('Acceptance: a new Google account reaches role selection with its real name and picture already '
  'populated, and profiles holds a matching row.', bold=True, color=GREEN)

# ==========================================================================
h1('7. Phase 12 — QR Check-in and Check-out')

p('The core of the requirement: "owner jkhn registration korbe tkhn 1ta unique qr code generate hbe, '
  'jeta owner download kore print korte parbe ... rider er qr code ta scan korte hbe park korar jonno."')

h2('7.1 Lifecycle')
code(
"""  Booked            Confirmed     (wallet already debited, interval reserved)
     |
     |  rider scans QR at the space  ->  check_in()
     v
  Parking           Active        (checked_in_at set, owner notified)
     |
     |  rider scans QR on exit      ->  check_out()
     v
  Settled           Completed     (checked_out_at, owner_earnings credited,
                                   interval released, review prompt queued)

  Divergences
  - scans out after end_time      -> overstay settled at exit, then Completed
  - never scans in                -> mark_no_shows() -> Cancelled (no_show)
  - scans in, never scans out     -> settle_stale_bookings() closes at end_time
                                     + accrued penalty, owner still paid""")

h2('7.2 check_in contract')
code(
"""check_in(p_token uuid, p_short_code text, p_booking uuid,
         p_lat numeric, p_lng numeric) RETURNS jsonb

Validation, in order — first failure returns and is logged to scan_events:
 1. caller is authenticated                      -> 'Not authenticated'
 2. booking exists and rider_id = auth.uid()     -> 'Booking not found'
 3. status IN ('Confirmed','Pending')            -> 'Booking is <status>'
 4. token or short code resolves to a listing    -> 'Invalid QR code'
 5. resolved listing = booking.listing_id        -> 'Wrong parking spot'
 6. now >= start_time - grace                    -> 'Too early, starts at <t>'
 7. now <= start_time + no_show_minutes          -> 'Booking window expired'
 8. distance(p_lat,p_lng, listing) <= 150 m      -> advisory warning only

On success (single transaction):
   bookings.status          = 'Active'
   bookings.checked_in_at   = now()
   bookings.check_in_method = 'qr' | 'code'
   INSERT scan_events(kind='check_in', ok=true)
   INSERT notifications(owner, type='check_in_alert')
   INSERT notifications(rider, type='check_in_confirmed')
RETURNS {ok, msg, booking_id, ends_at, listing_title}""")

h2('7.3 check_out contract')
code(
"""check_out(p_token uuid, p_booking uuid) RETURNS jsonb

 1-5. same identity and spot checks as check_in
 6.   status = 'Active' or 'Overstayed'          -> 'Not checked in'

Settlement:
   over_h := GREATEST(0, ceil(extract(epoch from now()-end_time)/3600))
   penalty := over_h * hourly_rate * overstay_penalty_multiplier
   if penalty > 0:
        if wallet_balance >= penalty: deduct + txn('overstay_charge')
        else: bookings.payment_due = penalty
              profiles.has_payment_due = TRUE
        bookings.overstay_amount = penalty
   bookings.actual_end_time = now()
   bookings.checked_out_at  = now()
   PERFORM complete_booking(p_booking)   -- credits owner, releases interval
   INSERT notifications(rider, 'checkout_success')
   INSERT notifications(owner, 'space_freed')
RETURNS {ok, msg, overstay_hours, overstay_charge, total_paid}""")
p('The longer the rider stays, the higher the charge — the penalty is linear in elapsed hours and '
  'is settled at the exit scan, so the amount is always final and explainable.')

h2('7.4 Overstay engine rework')
p('process_overstays() currently marks a booking Overstayed, charges once, and calls release_slot(). '
  'Under QR that is wrong: a rider still physically parked would have their capacity handed to '
  'someone else. Reworked behaviour:')
table(
    ['Condition', 'Old', 'New'],
    [
        ['end_time passed, checked in, not out', 'Charge once, release slot, mark Overstayed', 'Mark Overstayed, warn, keep the interval held, accrue on exit'],
        ['end_time passed, never checked in', 'Charge as overstay', 'Handled by mark_no_shows(): Cancelled, refund per policy, interval freed'],
        ['Charged twice by overlapping ticks', 'Prevented by SKIP LOCKED', 'Prevented by SKIP LOCKED and by settling only at check_out'],
    ],
    widths=[1.9, 2.1, 2.2],
)

h2('7.5 Frontend')
table(
    ['File', 'Action'],
    [
        ['pubspec.yaml', 'Add qr_flutter ^4.1.0, mobile_scanner ^5.2.3, printing ^5.13.4, path_provider'],
        ['owner/.../listing_qr_screen.dart', 'NEW — renders the QR with the Parkfinity mark, the short code, the listing title and address; Download PNG; Share; Print A4 signage'],
        ['owner/.../widgets/qr_poster.dart', 'NEW — printable poster: QR, short code, bilingual "Scan to start parking" instructions, spot name'],
        ['owner/.../my_listings_screen.dart', 'Add a QR action to every listing card'],
        ['owner/.../add_listing_screen.dart', 'On successful creation, route straight to the QR screen with a "print this and mount it" prompt'],
        ['rider/.../scan_screen.dart', 'NEW — camera preview, torch toggle, scan frame overlay, manual code entry sheet, success and failure states with haptics'],
        ['rider/.../active_parking_screen.dart', 'Primary action becomes Scan to Start or Scan to End depending on status; show elapsed time and a live overstay estimate once past end_time'],
        ['shared/data/repositories/scan_repository.dart', 'NEW — checkIn / checkOut RPC wrappers with typed results'],
        ['android/.../AndroidManifest.xml', 'CAMERA permission and a parkfinity:// intent filter'],
        ['ios/Runner/Info.plist', 'NSCameraUsageDescription in English and Bengali'],
    ],
    widths=[2.3, 3.9],
)
p('Acceptance: an owner prints a QR, a rider with a Confirmed booking scans it and the status turns '
  'Active with the owner notified; scanning a different listing\'s QR is refused with "Wrong parking '
  'spot"; scanning again on exit completes the booking and credits the owner.', bold=True, color=GREEN)

# ==========================================================================
h1('8. Phase 13 — Interval Availability and Rider-Chosen Time')

p('Closes the requirement: "Parking spot gula erokom dynamically khali hbe ar vora thakbe" and '
  '"rider can choose his/her time accordingly through the app."')

h2('8.1 The scenario matrix')
p('Every case below is a test in the Phase 13 suite. Listing L has slot_capacity {"Car": 7}.')
table(
    ['#', 'Scenario', 'Expected'],
    [
        ['1', 'R1 books Car 14:00-15:00', '6 Car slots free 14:00-15:00; 7 free at every other hour'],
        ['2', 'R2 books Car 15:00-16:00 after case 1', 'Accepted — no overlap'],
        ['3', 'R3 books Car 14:30-15:30 after case 1', 'Accepted — 5 free during the 14:30-15:00 overlap'],
        ['4', 'Seven riders book Car 14:00-15:00', 'Seventh accepted, eighth refused "No Car slots for this time"'],
        ['5', 'Eighth rider books Car 15:00-16:00 after case 4', 'Accepted — capacity fully recycles'],
        ['6', 'Motorcycle booked while all 7 Car slots are taken', 'Accepted if slot_capacity has a Motorcycle entry — types are independent'],
        ['7', 'Two riders submit for the last slot within milliseconds', 'Exactly one succeeds; the other gets a clean refusal, no partial debit'],
        ['8', 'Monthly booking 1-31 Aug', 'Blocks that slot for the whole month; the remaining 6 stay bookable hour by hour'],
        ['9', 'R1 extends 14:00-15:00 to 16:00', 'Allowed only if the slot is free 15:00-16:00; otherwise refused with the next free window offered'],
        ['10', 'R1 cancels', 'Interval freed immediately; refund per policy'],
        ['11', 'R1 checks out at 14:40', 'Interval freed at 14:40, not 15:00 — actual_end_time governs'],
        ['12', 'R1 overstays to 15:25', 'Interval stays held to 15:25; a 15:00 booking by R2 is refused at reservation time, so the conflict cannot arise'],
        ['13', 'Booking spans midnight 23:00-01:00', 'Single range, no day-boundary special case'],
        ['14', 'Booking outside availability_schedule', 'Refused "Owner is closed at that time"'],
        ['15', 'start >= end, or start in the past', 'Refused at validation, before any lock is taken'],
    ],
    widths=[0.3, 3.0, 2.9],
)

h2('8.2 reserve_interval')
code(
"""reserve_interval(p_listing uuid, p_vtype text, p_qty int,
                 p_start timestamptz, p_end timestamptz)
RETURNS TABLE(ok boolean, msg text, available int)

BEGIN
  IF p_start >= p_end THEN RETURN (false,'Invalid time range',0); END IF;
  IF p_start < now() - interval '5 minutes'
     THEN RETURN (false,'Start time is in the past',0); END IF;

  PERFORM pg_advisory_xact_lock(hashtextextended(p_listing::text, 0));

  SELECT * INTO l FROM listings WHERE id = p_listing FOR SHARE;
  IF NOT FOUND OR NOT l.is_active
     THEN RETURN (false,'Listing unavailable',0); END IF;
  IF NOT (l.slot_capacity ? p_vtype)
     THEN RETURN (false,'This spot does not accept '||p_vtype,0); END IF;
  IF NOT within_schedule(l.availability_schedule, p_start, p_end)
     THEN RETURN (false,'Owner is closed at that time',0); END IF;

  cap  := (l.slot_capacity->>p_vtype)::int;
  used := COALESCE((SELECT SUM(slot_qty) FROM bookings
                     WHERE listing_id = p_listing
                       AND vehicle_type = p_vtype
                       AND status IN ('Pending','Confirmed','Active','Overstayed')
                       AND tstzrange(start_time,
                             COALESCE(actual_end_time, end_time), '[)')
                           && tstzrange(p_start, p_end, '[)')), 0);
  free := cap - used;
  IF free < p_qty THEN RETURN (false,'Only '||free||' left for this time',free);
  END IF;
  RETURN (true,'ok',free - p_qty);
END""")
p('The lock is transaction-scoped, so the caller inserts the booking row in the same transaction; '
  'the lock releases on commit and the newly inserted row is visible to the next contender.')

h2('8.3 Frontend — time selection')
table(
    ['File', 'Action'],
    [
        ['rider/.../checkout_screen.dart', 'Rebuilt: duration-type segmented control; explicit start date and time picker; end derived from unit count with a manual override for the hourly case; live availability probe on every change; itemised price breakdown; a clear notice that the server price is authoritative'],
        ['rider/.../widgets/time_range_picker.dart', 'NEW — presets (1h, 2h, 4h, today, this week, this month) plus a custom range; blocks past times; enforces the owner schedule'],
        ['rider/.../widgets/availability_strip.dart', 'NEW — a 24-hour horizontal band for the chosen day showing free capacity per hour, so the rider can see at a glance when the space is busy'],
        ['rider/.../listing_details_screen.dart', 'Replace the static slot count with a date-aware availability strip'],
        ['rider/data/repositories/availability_repository.dart', 'NEW — listing_availability and available_between wrappers with a short-lived cache'],
        ['shared/.../bookings_repository.dart', 'BookingRequest gains explicit endTime; server response carries price_breakdown'],
    ],
    widths=[2.0, 4.2],
)

h2('8.4 Edge function changes')
bullet('create-booking: swap book_slot for reserve_interval; accept an explicit end_time; persist price_breakdown; release nothing on the happy path because reservation is now implied by the row itself.')
bullet('extend-booking: probe availability for [old_end, new_end) rather than consulting next_booking_start; on refusal, return the next free window so the UI can offer it.')
bullet('_shared/pricing.ts: apply the Dhaka offset (ADR-03); read peak_windows and weekend_days from settings; clamp the composed multiplier to max_price_multiplier; return every factor in the breakdown.')
p('Migration safety: existing bookings already carry start_time and end_time, so the interval query '
  'is correct for historical rows on day one. slot_available is recomputed once at migration time '
  'and thereafter by trigger.', italic=True)

# ==========================================================================
h1('9. Phase 14 — Booking Details, Contact, Navigation')

p('Closes D-07 and D-08 and the click-to-call requirement.')

h2('9.1 Shared detail screen')
table(
    ['Section', 'Rider view', 'Owner view'],
    [
        ['Header', 'Listing title, photo, status chip', 'Rider name, avatar, status chip'],
        ['Counterparty', 'Owner name, rating, Call and Message', 'Rider name, rating, vehicle, Call'],
        ['Timing', 'Start, end, duration, checked-in and checked-out stamps, live countdown', 'Same'],
        ['Vehicle', 'Type, plate, brand, colour', 'Same — the owner needs the plate to identify the car'],
        ['Location', 'Address, static map thumbnail, Navigate', 'Address only'],
        ['Money', 'Base, peak, weekend, demand, total, commission shown as "platform fee", overstay if any', 'Gross, platform commission, net earnings, payout status'],
        ['Access', 'Scan to Start / Scan to End, or the check-in timeline once complete', 'Scan history from scan_events'],
        ['Actions', 'Extend, Cancel, Rate, Report a problem', 'Approve or Decline for manual mode, Rate the rider, Report'],
    ],
    widths=[1.0, 2.6, 2.6],
)

h2('9.2 Click-to-call')
code(
"""Future<void> callParty(String rawPhone) async {
  final uri = Uri(scheme: 'tel', path: normalizeBd(rawPhone));
  if (!await launchUrl(uri)) { /* localized "no dialer" message */ }
}""")
p('launchUrl with a tel: scheme opens the dialer pre-filled; it never places the call without the '
  'user pressing dial, which is the correct behaviour on both platforms.')
p('Privacy rule: the counterparty phone number is returned by the API only while a booking between '
  'the two parties is Confirmed, Active, or Overstayed, and for 24 hours after completion. Outside '
  'that window the field is null and the Call button is hidden. This is enforced in the '
  'booking_detail() RPC, not in the UI.', bold=True)

h2('9.3 Navigation hand-off')
code(
"""google.navigation:q=<lat>,<lng>&mode=d      // Android, turn-by-turn
comgooglemaps://?daddr=<lat>,<lng>&directionsmode=driving   // iOS
https://www.google.com/maps/dir/?api=1&destination=<lat>,<lng>  // fallback""")
p('The first URI that canLaunchUrl accepts is used, so the rider lands in the Google Maps app '
  'wherever it is installed and in the browser otherwise.')

h2('9.4 Files')
table(
    ['File', 'Action'],
    [
        ['shared/presentation/screens/booking_detail_screen.dart', 'NEW — role-aware, ~600 lines, composed from section widgets'],
        ['shared/presentation/widgets/booking_status_chip.dart', 'NEW — one status vocabulary for every screen'],
        ['shared/domain/booking_status_policy.dart', 'NEW — classification logic extracted from three screens (D-08)'],
        ['shared/presentation/widgets/contact_row.dart', 'NEW — avatar, name, rating, call button'],
        ['shared/presentation/widgets/price_breakdown_card.dart', 'NEW — itemised, localized, currency-formatted'],
        ['rider/.../rider_booking_history_screen.dart', 'Wrap cards in InkWell; fix tab filters via the policy; add Cancelled and Refunded chips'],
        ['owner/.../owner_booking_history_screen.dart', 'Same treatment; add Approve and Decline for manual-mode listings'],
        ['core/router/app_router.dart', 'Add /booking/:id under both shells'],
    ],
    widths=[2.6, 3.6],
)

# ==========================================================================
h1('10. Phase 15 — Wallet Role Split and Payout Automation')

p('Closes D-09 and the unfinished payout path.')

h2('10.1 Screen split')
table(
    ['Screen', 'Route', 'Contents'],
    [
        ['RiderWalletScreen', '/rider/wallet', 'Balance, Add Funds via SSLCommerz, spend history, refunds, outstanding payment_due banner with a Pay Now action'],
        ['OwnerEarningsScreen', '/owner/wallet', 'Available balance, pending clearance, lifetime gross, commission paid, Withdraw, earnings history, a 30-day trend chart — no top-up anywhere'],
    ],
    widths=[1.5, 1.3, 3.4],
)

h2('10.2 Enforcement')
code(
"""-- UI cannot be the only guard.
ALTER TABLE transactions ADD CONSTRAINT no_owner_deposit CHECK (
  type <> 'deposit'
  OR NOT EXISTS (SELECT 1 FROM profiles p
                  WHERE p.id = transactions.user_id AND p.role = 'Owner')
);""")
p('Implemented as a BEFORE INSERT trigger rather than a literal CHECK, since a CHECK constraint '
  'cannot contain a subquery. Same effect, enforced in the database.', italic=True, color=GREY)

h2('10.3 owner_earnings_summary')
code(
"""owner_earnings_summary(p_owner uuid) RETURNS jsonb
{
  "available":        profiles.wallet_balance - SUM(pending withdrawal holds),
  "pending_clearance": SUM(owner_earnings) WHERE status IN
                       ('Confirmed','Active','Overstayed'),
  "lifetime_gross":    SUM(total_amount)   WHERE status = 'Completed',
  "lifetime_net":      SUM(owner_earnings) WHERE status = 'Completed',
  "commission_paid":   SUM(commission_amount) WHERE status = 'Completed',
  "withdrawn":         SUM(amount) FROM withdrawals WHERE status = 'Completed',
  "next_payout_eta":   MIN(end_time) WHERE status IN ('Confirmed','Active')
}""")
p('"Pending clearance" is the figure that answers the owner question directly: money already '
  'earned but not yet released, because the rider has not checked out.')

h2('10.4 Closing the payout gap')
p('Three independent paths now credit the owner, so no booking can strand:')
num('QR check-out — the normal path. complete_booking() runs inside check_out().')
num('settle_stale_bookings() — pg_cron hourly. Any booking checked in but not out, more than two hours past end_time, is closed at end_time plus accrued penalty and the owner is paid.')
num('Admin force-complete — a manual action in the admin panel for dispute resolution, written to audit_log.')
p('Acceptance: an owner who never touches the app sees earnings appear within an hour of a rider '
  'finishing, and the Add Funds affordance does not exist on any owner surface.', bold=True, color=GREEN)

# ==========================================================================
h1('11. Phase 16 — Explore Map Upgrade')

p('Closes rider items 1 and 7 and D-12.')

h2('11.1 Search with suggestions')
p('Reported: "if I write any location it would suggest me the locations of the same names from '
  'where I can choose my desired location." Today _searchPlace geocodes only on submit and takes '
  'the first hit blindly.')
bullet('A debounced 300 ms query hits the places-autocomplete edge function (ADR-05).')
bullet('Results render in an overlay list under the search bar: bold primary text, grey secondary text, a distance badge.')
bullet('Selecting a suggestion resolves coordinates through Place Details, animates the camera, and drops a temporary search-target pin.')
bullet('Recent searches persist in shared_preferences and appear when the field is focused but empty.')
bullet('If the edge function is unreachable, the existing geocoding path is used silently as a fallback.')

h2('11.2 Automatic GPS')
p('Reported: "the gps will automatically work, I don\'t have to press the button." '
  '_checkLocationPermission already runs in initState, but _goToCurrentLocation awaits '
  '_controller.future, which only completes in onMapCreated — so the first fix is frequently '
  'dropped and the map stays on the Dhaka fallback.')
bullet('Resolve the position and the map controller independently, then apply the camera move when both are ready.')
bullet('Subscribe to Geolocator.getPositionStream with a 25 m distance filter so the rider marker tracks movement.')
bullet('Show an explicit rationale sheet when permission is denied, with a Settings deep link.')
bullet('The manual button is retained as a re-centre control, not as the only way to get a fix.')

h2('11.3 Marker vocabulary')
table(
    ['Marker', 'Appearance', 'Meaning'],
    [
        ['Rider', 'Blue pulsing dot with a heading cone (custom bitmap)', 'You are here'],
        ['Available listing', 'Purple pin with the hourly price in the label', 'Bookable now'],
        ['Full listing', 'Grey pin at 50% opacity', 'Exists, no capacity for the chosen window'],
        ['AI top pick', 'Gold pin with a sparkle badge', 'Highest recommendation score'],
        ['Search target', 'Red teardrop, dismissible', 'Where you searched'],
        ['Cluster', 'Purple circle with a count', 'More than 12 pins within 60 px'],
    ],
    widths=[1.1, 2.6, 2.5],
)
p('The current implementation uses hueViolet for available and hueRose for full, and relies on '
  'the built-in blue dot for the rider — which is easy to lose among the pins. Custom bitmaps '
  'render from Canvas at device pixel ratio so labels stay crisp.')

h2('11.4 Recommendations use the live position')
p('The scoring engine already weights distance 0.35, price 0.25, rating 0.20, security 0.10, and '
  'history 0.10. Changes:')
bullet('Feed the streamed position rather than a snapshot, so the ranking updates as the rider moves.')
bullet('Score only listings that have capacity for the rider\'s selected time window — a cheap, close, unavailable spot is not a recommendation.')
bullet('Add an availability term so a spot with 6 of 7 free outranks one with 1 of 7.')
bullet('Show the top three as a swipeable carousel over the map, each card tappable straight through to details.')

h2('11.5 Photo gallery')
bullet('PhotoGalleryScreen: PageView plus InteractiveViewer, pinch and double-tap zoom, a "3 / 8" counter, swipe-down to dismiss, Hero transition from the tapped thumbnail.')
bullet('Thumbnail strip under the hero image on the details screen; tapping any thumbnail opens the gallery at that index.')
bullet('Video, where the listing has one, appears as the first gallery entry with a play overlay.')

# ==========================================================================
h1('12. Phase 17 — Theme, Localization, Static Pages')

p('Closes D-10, D-11, D-13.')

h2('12.1 Dark mode remediation')
p('Substitution table applied across the 23 affected files:')
table(
    ['Hardcoded', 'Replacement'],
    [
        ['Colors.white (surface)', 'Theme.of(context).colorScheme.surface'],
        ['Colors.grey[50] (scaffold)', 'Theme.of(context).colorScheme.surfaceContainerLowest'],
        ['Colors.grey[600] (caption)', 'Theme.of(context).colorScheme.onSurfaceVariant'],
        ['Colors.grey.shade200 (border)', 'Theme.of(context).colorScheme.outlineVariant'],
        ['Colors.deepPurple (brand)', 'Theme.of(context).colorScheme.primary'],
        ['Colors.black (body text)', 'Theme.of(context).colorScheme.onSurface'],
        ['Colors.red (destructive)', 'Theme.of(context).colorScheme.error'],
        ['Semantic status colours', 'A ParkfinityColors ThemeExtension with light and dark variants for success, warning, info, and pending'],
    ],
    widths=[2.4, 3.8],
)
p('Both ThemeData objects gain explicit scaffoldBackgroundColor, cardTheme, dividerTheme, '
  'listTileTheme, chipTheme, dialogTheme, and bottomSheetTheme so components stop falling back to '
  'Material defaults. A ThemeExtension carries the brand semantics that ColorScheme has no slot for.')
p('Guard: an analyzer rule plus a CI grep that fails the build if Colors.white, Colors.black, or '
  'Colors.grey[ appears in lib/features. Existing exceptions are annotated inline.', bold=True)

h2('12.2 Bengali completion')
num('Extract the ~70 remaining literals into app_en.arb with descriptive keys and ICU placeholders.')
num('Translate every key into app_bn.arb, reaching parity at roughly 360 keys.')
num('Localize dates (DateFormat with the active locale), currency (Taka with Bengali numerals when the locale is bn), and relative times ("2 hours ago" / "২ ঘণ্টা আগে").')
num('Localize server-produced strings — booking status names, notification bodies, and RPC messages — by returning a message code plus parameters rather than English prose, and resolving it client-side.')
num('Add test/l10n_parity_test.dart asserting identical key sets, and run flutter gen-l10n before every analyze (a known trap from Phase 10: generated Dart goes stale against the ARB files and surfaces as undefined_getter).')
p('Bengali numerals: ০১২৩৪৫৬৭৮৯ are used for prices and counts under the bn locale via '
  'NumberFormat with the bn_BD symbol set.')

h2('12.3 Help Center and Privacy Policy')
p('Content is stored in the static_pages table (slug, locale, title, body markdown, updated_at) so '
  'the admin can edit it without an app release, with the seeded copy bundled as an offline fallback.')
table(
    ['Page', 'Sections'],
    [
        ['Help Center', 'Getting started; How booking works; Scanning the QR code; Payments and refunds; Overstay charges; For space owners; Withdrawals; Safety; Contact support'],
        ['Privacy Policy', 'What we collect; Location data; Documents and KYC; Payment data; How we use it; Sharing with owners and riders; Retention; Your rights; Cookies; Changes; Contact'],
    ],
    widths=[1.2, 5.0],
)
p('Both are supplied as placeholder copy in English and Bengali, clearly marked as a template '
  'requiring legal review before public release.', italic=True, color=GREY)

# ==========================================================================
h1('13. Phase 18 — Notification Completeness')

p('Closes the requirement: "Sob dhoroner notification jno thik vabe handle hoy."')

h2('13.1 Full matrix')
p('Producer column names what emits the row. Every row inserted into notifications is pushed '
  'automatically by the Phase 8 dispatch_push trigger.')
table(
    ['Type', 'Recipient', 'Trigger point', 'Producer', 'Status'],
    [
        ['booking_confirmed', 'Rider', 'Booking created, instant mode', 'create-booking', 'exists'],
        ['booking_requested', 'Owner', 'Booking created, manual mode', 'create-booking', 'exists'],
        ['booking_approved', 'Rider', 'Owner approves a manual request', 'approve_booking', 'NEW'],
        ['booking_declined', 'Rider', 'Owner declines; full refund', 'decline_booking', 'NEW'],
        ['new_booking_alert', 'Owner', 'Any booking on their listing', 'create-booking', 'exists'],
        ['payment_success', 'Rider', 'Wallet debit or SSLCommerz settled', 'create-booking / webhook', 'exists'],
        ['payment_failed', 'Rider', 'Insufficient balance or gateway failure', 'create-booking', 'NEW'],
        ['wallet_topup', 'Rider', 'Funds added', 'sslcommerz-webhook', 'NEW'],
        ['check_in_alert', 'Owner', 'Rider scans in', 'check_in', 'NEW'],
        ['check_in_confirmed', 'Rider', 'Scan accepted', 'check_in', 'NEW'],
        ['booking_reminder', 'Rider', 'reminder_minutes before end_time', 'notification_sweep', 'NEW'],
        ['start_reminder', 'Rider', '30 min before start_time', 'notification_sweep', 'NEW'],
        ['overstay_warning', 'Rider', 'At end_time, still checked in', 'notification_sweep', 'NEW'],
        ['overstay_charged', 'Rider', 'Penalty applied', 'check_out / process_overstays', 'exists'],
        ['overstay_alert', 'Owner', 'Their space is held past its booking', 'notification_sweep', 'NEW'],
        ['checkout_success', 'Rider', 'Scanned out', 'check_out', 'NEW'],
        ['space_freed', 'Owner', 'Rider left', 'check_out', 'NEW'],
        ['no_show', 'Both', 'Never scanned in within the window', 'mark_no_shows', 'NEW'],
        ['booking_cancelled', 'Both', 'Either party cancels', 'cancel_booking', 'exists'],
        ['refund_issued', 'Rider', 'Refund credited', 'cancel_booking', 'exists'],
        ['booking_extended', 'Both', 'Extension accepted', 'extend-booking', 'NEW'],
        ['earnings_credited', 'Owner', 'complete_booking pays out', 'complete_booking', 'exists'],
        ['withdrawal_requested', 'Owner', 'Request submitted', 'request_withdrawal', 'NEW'],
        ['withdrawal_approved', 'Owner', 'Admin approves', 'approve_withdrawal', 'NEW'],
        ['withdrawal_rejected', 'Owner', 'Admin rejects; funds released', 'reject_withdrawal', 'NEW'],
        ['review_prompt', 'Both', '1 hour after completion, if unreviewed', 'notification_sweep', 'NEW'],
        ['review_received', 'Both', 'Counterparty submits a review', 'submit_review', 'NEW'],
        ['kyc_verified', 'User', 'Admin verifies documents', 'admin_set_kyc', 'NEW'],
        ['kyc_rejected', 'User', 'Admin rejects documents', 'admin_set_kyc', 'NEW'],
        ['listing_approved', 'Owner', 'Admin approves a listing', 'admin action', 'NEW'],
        ['account_suspended', 'User', 'Admin suspends', 'admin_set_suspended', 'NEW'],
        ['payment_due', 'Rider', 'Unsettled overstay blocks new bookings', 'process_overstays', 'exists'],
    ],
    widths=[1.4, 0.8, 1.9, 1.4, 0.7],
)

h2('13.2 notification_sweep')
code(
"""notification_sweep() -- pg_cron: '*/5 * * * *'
  1. start_reminder    : status='Confirmed'
                         AND start_time BETWEEN now() AND now()+30 min
  2. booking_reminder  : status IN ('Confirmed','Active')
                         AND end_time BETWEEN now()
                             AND now() + reminder_minutes
  3. overstay_warning  : status='Active' AND end_time < now()
  4. overstay_alert    : same rows, addressed to the owner
  5. review_prompt     : status='Completed'
                         AND updated_at < now()-1 hour
                         AND no review row for that side
  Every insert is ON CONFLICT DO NOTHING against idx_notif_once,
  so overlapping ticks cannot duplicate.""")

h2('13.3 Delivery quality')
bullet('Per-user preferences in profiles.notification_prefs: a channel switch (push, in-app, SMS) and a per-type opt-out. Transactional types — payment, overstay, cancellation — cannot be disabled.')
bullet('Notification bodies are stored as a message code plus a JSON parameter bag, so the client renders them in the active language instead of receiving fixed English text.')
bullet('Deep links: every notification carries a route in its data payload, so tapping a booking reminder opens that booking, not the home screen.')
bullet('Grouping: Android channels per category (bookings, payments, alerts) with distinct importance levels; alerts bypass Do Not Disturb, marketing does not.')
bullet('SMS remains a stub interface with a single implementation point, ready for a provider once one is procured — the requirement lists it, but no gateway is configured.')

# ==========================================================================
h1('14. Phase 19 — Hardening, Admin Catch-up, QA')

h2('14.1 Correctness and abuse')
table(
    ['Item', 'Measure'],
    [
        ['Duplicate submissions', 'bookings.idempotency_key, client-generated UUID per checkout attempt; a repeat returns the original booking rather than creating a second'],
        ['Race on last slot', 'Advisory lock in reserve_interval (ADR-02), covered by a concurrent test'],
        ['Wallet drift', 'A nightly reconciliation job asserting wallet_balance equals the signed sum of that user\'s transactions; discrepancies raise an admin alert'],
        ['Booking spam', 'Rate limit of 10 booking attempts per rider per hour, enforced in create-booking against a counter table'],
        ['QR replay', 'scan_events records every attempt; more than 5 failures on one booking within 10 minutes locks scanning for that booking and notifies both parties'],
        ['Token leakage', 'listings.qr_token is never exposed by any client-readable view; check_in resolves it server-side only'],
        ['Suspended accounts', 'Already checked in create-booking; extended to check_in, extend-booking, and withdrawal'],
        ['Payment due', 'Already blocks new bookings; a Pay Now action is added so riders can clear it'],
    ],
    widths=[1.4, 4.8],
)

h2('14.2 Admin panel catch-up')
bullet('QR management: view, regenerate, and download any listing QR; regeneration writes qr_rotated_at and notifies the owner to reprint.')
bullet('Scan log viewer per booking, with the rejection reason for failures.')
bullet('Force check-out and force-complete for dispute resolution, both written to audit_log.')
bullet('Static page editor for Help Center and Privacy Policy in both languages.')
bullet('Notification preferences and a manual broadcast to a filtered user segment.')
bullet('Peak window and weekend day editor exposed in the settings screen alongside the existing six multipliers.')
bullet('No-show report and an overstay report with revenue attribution.')

h2('14.3 Test plan')
table(
    ['Level', 'Coverage'],
    [
        ['Database', 'pgTAP over reserve_interval (all 15 scenarios in 8.1), check_in and check_out validation ladders, cancel and refund policy boundaries, commission arithmetic to two decimal places'],
        ['Concurrency', 'Twenty parallel reserve_interval calls for the last slot; exactly one must succeed and no wallet may be debited on a failure'],
        ['Edge functions', 'Deno tests for pricing (timezone boundaries at 07:59, 08:00, 09:59, 10:00 Dhaka; Thursday-to-Sunday weekend edges), create-booking rejection paths, extend-booking guards'],
        ['Widget', 'Time range picker, availability strip, price breakdown card, booking status chip, avatar fallback'],
        ['Integration', 'Register to book to scan in to scan out to review, for both roles, in both languages, in both themes'],
        ['Localization', 'Key parity; a golden test that no screen renders a raw ARB key'],
        ['Theme', 'Golden tests for eight key screens in light and dark; a contrast assertion at WCAG AA'],
        ['Manual', 'Physical QR print, scanned from a phone at realistic distance and lighting'],
    ],
    widths=[1.0, 5.2],
)

h2('14.4 Manual steps still outstanding from Phase 8')
p('Push notifications remain undelivered until both of these are done with real credentials. They '
  'cannot be committed to the repository.', color=RED, bold=True)
code(
"""1. supabase secrets set FCM_SERVICE_ACCOUNT="$(cat service-account.json)"
   # Firebase console -> Project settings -> Service accounts -> Generate key

2. UPDATE app_config
      SET functions_url   = 'https://rkqduzjkkyplceipydir.functions.supabase.co',
          service_role_key = '<service-role-key>'
    WHERE id = TRUE;
   # Supabase dashboard -> Project settings -> API -> service_role""")
p('Until both are set the dispatch trigger skips silently and in-app notifications still save '
  'correctly, so nothing breaks — the messages simply do not reach the device.')

# ==========================================================================
h1('15. Additional Recommendations')

p('Beyond the stated scope. Each is independently shippable; none is assumed.')

h2('15.1 Product')
table(
    ['Idea', 'Rationale', 'Effort'],
    [
        ['Favourites / saved spots', 'Commuters park in the same place daily; one tap to rebook', 'S'],
        ['Rebook last booking', 'The single highest-frequency action in a commuter marketplace', 'S'],
        ['Recurring bookings', 'Weekday 09:00-18:00 for a month, generated as linked bookings', 'M'],
        ['Waitlist on full spots', 'Notify when capacity frees in a watched window; captures demand that currently bounces', 'M'],
        ['Owner calendar view', 'A month grid of bookings per listing — the clearest occupancy signal an owner can get', 'M'],
        ['Dispute flow', 'Structured report with photos, feeding the admin force-complete path', 'M'],
        ['Referral credit', 'Both sides of a marketplace need supply; owner referrals compound', 'S'],
        ['Listing analytics', 'Views, conversion, revenue per slot — helps owners price correctly', 'M'],
    ],
    widths=[1.5, 3.7, 0.5],
)

h2('15.2 Technical')
bullet('Move latitude and longitude to PostGIS geography and use ST_DWithin, replacing the client-side Haversine sweep. Radius search stops being O(all listings) and becomes index-backed.')
bullet('Add a listings_search materialised view with a tsvector over title, address, and area for text search alongside map search.')
bullet('Introduce Sentry (or Supabase logging) for client crash reporting — currently a production crash is invisible.')
bullet('Add a feature-flag table so risky changes such as interval availability can be dark-launched per user.')
bullet('Cache listing photos with cached_network_image; the map and details screens currently refetch on every build.')
bullet('Extract a domain layer of pure Dart entities and use-cases. The current repositories return data models straight to widgets, which is workable but couples presentation to the wire format — the stated Clean Architecture goal is only partially met.')
bullet('Add supabase db diff to CI so schema drift between the repository and the remote project is caught before it reaches a release.')

h2('15.3 Business logic worth deciding')
table(
    ['Question', 'Recommendation'],
    [
        ['No-show refund policy', 'Charge 25% as a cancellation fee, refund 75%, pay the owner the fee — the space was genuinely held'],
        ['Overstay ceiling', 'Cap accrual at 24 hours of penalty, then escalate to the owner as a dispute rather than accruing indefinitely'],
        ['Minimum payout', 'Set a floor of 500 Taka per withdrawal to control transfer costs'],
        ['Commission tiers', 'Reduce the rate for high-volume owners, e.g. 15% down to 10% above 50 bookings per month, as a retention lever'],
        ['Cancellation window', 'Free cancellation up to 2 hours before start; the existing refund_full_hours setting already expresses this'],
        ['Owner response SLA', 'Auto-decline manual requests unanswered after 30 minutes and refund in full, so riders are never left waiting'],
    ],
    widths=[1.8, 4.4],
)

# ==========================================================================
h1('16. Execution Order and Verification')

h2('16.1 Sequence')
table(
    ['Step', 'Work', 'Gate before moving on'],
    [
        ['1', 'Phase 11 backend migration and RPCs', 'Migration applied; my_profile and update_my_profile return correctly for a live user'],
        ['2', 'Phase 11 frontend', 'Sign-up completes; Google sign-in reaches role selection; the profile shows the real identity; edit and avatar persist'],
        ['3', 'Phase 13 backend (before 12 — check-out depends on it)', 'All 15 availability scenarios pass; the concurrency test shows exactly one winner'],
        ['4', 'Phase 13 edge functions and frontend', 'A booking with an explicit time range prices correctly with a visible breakdown'],
        ['5', 'Phase 12 backend', 'check_in and check_out ladders pass; the overstay rework does not release held capacity'],
        ['6', 'Phase 12 frontend', 'A printed QR scans and starts a booking; a second scan finishes it and pays the owner'],
        ['7', 'Phase 14', 'Both roles open booking details; call and navigate work on a device'],
        ['8', 'Phase 15', 'Owners have no top-up affordance; earnings appear automatically'],
        ['9', 'Phase 16', 'Suggestions appear while typing; GPS resolves unaided; the gallery opens'],
        ['10', 'Phase 17', 'Dark mode transforms every screen; Bengali is complete; both static pages render'],
        ['11', 'Phase 18', 'Every matrix row has a producer and delivers'],
        ['12', 'Phase 19', 'Full suite green; flutter analyze reports zero errors and zero warnings'],
    ],
    widths=[0.4, 2.5, 3.3],
)
p('Phase 13 precedes Phase 12 deliberately: check_out must release an interval, and that concept '
  'does not exist until 13 lands. Building 12 first would mean writing release logic twice.',
  italic=True)

h2('16.2 Deployment method')
p('Established in Phase 0 and unchanged. No Docker, no psql, no database password.')
code(
"""supabase db push --linked            # migrations (the Docker warning is harmless)
supabase functions deploy <name> --no-verify-jwt
supabase secrets set KEY=value
# Remote schema introspection via PostREST with the anon key;
# a missing column surfaces as error 42703.""")

h2('16.3 Definition of done')
bullet('Every defect D-01 to D-13 verified fixed on a device, not merely in code.')
bullet('flutter analyze: zero errors, zero warnings, in both parkfinity3 and parkfinity_admin.')
bullet('flutter gen-l10n run and committed; the parity test passes.')
bullet('Every new RPC exercised live against ParkFinityDB and the result recorded.')
bullet('The full rider journey completed end to end in Bengali and in dark mode.')
bullet('The full owner journey completed including printing a QR and receiving a payout.')
bullet('No secret, key, or service account present in any committed file.')

doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run('— End of Implementation Plan v2 —')
r.font.color.rgb = GREY
r.italic = True

os.makedirs(os.path.dirname(OUT), exist_ok=True) if os.path.dirname(OUT) else None
doc.save(OUT)
print('WROTE', OUT)
