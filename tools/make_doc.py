import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(14.5)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(12.5)
    r.font.color.rgb = RGBColor(0x2B, 0x4C, 0x7E)
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_body_p(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def add_bullet(doc, bold_prefix, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.25)
    
    r_bullet = p.add_run("•  ")
    r_bullet.font.name = "Times New Roman"
    r_bullet.font.size = Pt(11)
    r_bullet.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.bold = True
        r_pre.font.name = "Times New Roman"
        r_pre.font.size = Pt(11)
        r_pre.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        
    r_txt = p.add_run(text)
    r_txt.font.name = "Times New Roman"
    r_txt.font.size = Pt(11)
    r_txt.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def add_step(doc, step_num, step_title, step_desc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Inches(0.25)
    
    r_num = p.add_run(f"Step {step_num}: ")
    r_num.bold = True
    r_num.font.name = "Times New Roman"
    r_num.font.size = Pt(11)
    r_num.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    if step_title:
        r_title = p.add_run(f"{step_title} – ")
        r_title.bold = True
        r_title.font.name = "Times New Roman"
        r_title.font.size = Pt(11)
        r_title.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
    r_desc = p.add_run(step_desc)
    r_desc.font.name = "Times New Roman"
    r_desc.font.size = Pt(11)
    r_desc.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    return p

def add_fig(doc, img_path, fig_title, fig_desc, is_mobile=True):
    if not os.path.exists(img_path):
        print(f"Warning: Image missing: {img_path}")
        return
    
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(8)
    p_img.paragraph_format.space_after = Pt(3)
    p_img.paragraph_format.keep_with_next = True
    
    run_img = p_img.add_run()
    if is_mobile:
        run_img.add_picture(img_path, width=Inches(2.3))
    else:
        run_img.add_picture(img_path, width=Inches(5.7))
        
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_before = Pt(2)
    p_cap.paragraph_format.space_after = Pt(1)
    p_cap.paragraph_format.keep_with_next = True
    
    r_cap = p_cap.add_run(fig_title)
    r_cap.bold = True
    r_cap.font.name = "Times New Roman"
    r_cap.font.size = Pt(10)
    r_cap.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    p_exp = doc.add_paragraph()
    p_exp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_exp.paragraph_format.space_before = Pt(0)
    p_exp.paragraph_format.space_after = Pt(8)
    
    r_exp = p_exp.add_run(fig_desc)
    r_exp.italic = True
    r_exp.font.name = "Times New Roman"
    r_exp.font.size = Pt(9.5)
    r_exp.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def add_callout(doc, title, text, icon="ℹ️", box_color="1E3A8A", bg_color="F4F6F9"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    set_cell_background(cell, bg_color)
    
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{box_color}"/>'
        f'  <w:top w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    r_icon = p.add_run(f"{icon}  ")
    r_icon.font.name = "Segoe UI Emoji"
    r_icon.font.size = Pt(10)
    
    r_title = p.add_run(f"{title}: ")
    r_title.bold = True
    r_title.font.name = "Times New Roman"
    r_title.font.size = Pt(10)
    r_title.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    
    r_text = p.add_run(text)
    r_text.font.name = "Times New Roman"
    r_text.font.size = Pt(10)
    r_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    p_sp = doc.add_paragraph()
    p_sp.paragraph_format.space_before = Pt(0)
    p_sp.paragraph_format.space_after = Pt(4)

print("make_doc helpers ready.")
